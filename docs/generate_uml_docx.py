"""
Genera el documento Word con los 5 diagramas UML de ms-auditoria.
Ejecutar: python docs/generate_uml_docx.py
Requiere: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
    return heading

def add_code_block(doc, text):
    """Agrega un bloque de texto monoespaciado (para diagramas ASCII)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p

def add_table(doc, headers, rows, col_widths=None):
    """Crea tabla con encabezados y filas."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)

    # Rows
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    return table


def generate_docx():
    doc = Document()

    # ── Configuración general ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    sections = doc.sections
    for section in sections:
        section.page_width = Cm(29.7)  # A4 landscape
        section.page_height = Cm(21.0)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # ══════════════════════════════════════════════════════════════
    # PORTADA
    # ══════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_heading("Diagramas UML — ms-auditoria", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Microservicio #19: Auditoría y Logging del ERP Universitario")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    info_lines = [
        "Materia: Desarrollo de Software 3",
        "Fecha: Marzo 2026",
        "Repositorio: https://github.com/Karatsuyu/ms-auditoria",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # TABLA DE CONTENIDOS
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Tabla de Contenidos", 1)
    toc_items = [
        "1. Diagrama de Casos de Uso",
        "   1.1 Actores del Sistema",
        "   1.2 Diagrama de Casos de Uso",
        "   1.3 Descripción de Casos de Uso",
        "2. Diagrama de Clases UML",
        "   2.1 Diagrama Completo",
        "   2.2 Relaciones entre Clases",
        "   2.3 Tabla de Relaciones",
        "3. Diagramas de Secuencia UML",
        "   3.1 Crear Log de Auditoría (POST /log)",
        "   3.2 Consultar Logs con Filtros (GET /logs)",
        "   3.3 Scheduler de Retención Automática (TTL)",
        "   3.4 Validación de API Key (Detalle)",
        "   3.5 Obtener Estadísticas (GET /stats)",
        "4. Diagrama de Componentes Interno",
        "   4.1 Diagrama de Componentes (Clean Architecture)",
        "   4.2 Diagrama Simplificado de Capas MVC",
        "5. Diagrama Entidad-Relación (ER)",
        "   5.1 Diagrama ER Completo",
        "   5.2 Notación ER con Cardinalidades",
        "   5.3 Detalle de Atributos con Tipos",
        "   5.4 Resumen de Índices",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 1. DIAGRAMA DE CASOS DE USO
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "1. Diagrama de Casos de Uso", 1)

    # 1.1 Actores
    add_heading_styled(doc, "1.1 Actores del Sistema", 2)

    actores_headers = ["Actor", "Tipo", "Descripción"]
    actores_rows = [
        ["Microservicio Externo", "Sistema externo", "Cualquiera de los 18 microservicios del ERP que envían eventos de auditoría"],
        ["Administrador del Sistema", "Usuario humano", "Persona que consulta logs, revisa estadísticas, gestiona tokens y ejecuta purgas manuales"],
        ["Scheduler (RetentionService)", "Sistema interno", "Proceso en background que purga logs antiguos según la política TTL"],
        ["ms-autenticación", "Sistema externo", "Microservicio que valida tokens de sesión de usuarios"],
        ["ms-roles-permisos", "Sistema externo", "Microservicio que verifica permisos de usuarios"],
    ]
    add_table(doc, actores_headers, actores_rows)
    doc.add_paragraph()

    # 1.2 Diagrama visual
    add_heading_styled(doc, "1.2 Diagrama de Casos de Uso", 2)

    cu_diagram = """
╔═══════════════════════════════════════════════════════════════════════════╗
║                          ms-auditoria (Sistema)                         ║
║                                                                         ║
║  ┌──────────────────────── GESTIÓN DE LOGS ─────────────────────────┐   ║
║  │                                                                   │   ║
║  │  (CU-01) Registrar Log         (CU-02) Registrar Logs en Batch   │   ║
║  │        │ «include»                    │ «include»                 │   ║
║  │        └────────► (CU-03) Validar API Key ◄───────┘              │   ║
║  │                                                                   │   ║
║  └───────────────────────────────────────────────────────────────────┘   ║
║                                                                         ║
║  ┌──────────────────────── CONSULTA DE LOGS ────────────────────────┐   ║
║  │                                                                   │   ║
║  │  (CU-04) Consultar Logs      (CU-05) Filtrar Logs Avanzado      │   ║
║  │  (CU-06) Obtener por ID      (CU-07) Trazar Request (X-Req-ID)  │   ║
║  │  (CU-08) Logs por Usuario    (CU-09) Buscar Full-Text en Detalle│   ║
║  │                                                                   │   ║
║  └───────────────────────────────────────────────────────────────────┘   ║
║                                                                         ║
║  ┌──────────────────────── ESTADÍSTICAS ────────────────────────────┐   ║
║  │  (CU-10) Obtener Estadísticas (logs/servicio, errores, avg dur) │   ║
║  └───────────────────────────────────────────────────────────────────┘   ║
║                                                                         ║
║  ┌──────────────────────── ADMINISTRACIÓN ──────────────────────────┐   ║
║  │                                                                   │   ║
║  │  (CU-11) Purgar Logs Antiguos (Manual)  ──«include»──► CU-03    │   ║
║  │  (CU-12) Purga Automática (TTL Scheduler)                       │   ║
║  │  (CU-13) Verificar Estado    (CU-14) Health Check               │   ║
║  │                                                                   │   ║
║  └───────────────────────────────────────────────────────────────────┘   ║
║                                                                         ║
╚═══════════════════════════════════════════════════════════════════════════╝

ACTORES:
  [Microservicio Externo] ──────► CU-01, CU-02 (vía CU-03)
  [Administrador]         ──────► CU-04..CU-11, CU-13, CU-14
  [Scheduler/Retention]   ──────► CU-12 (automático)
"""
    add_code_block(doc, cu_diagram)
    doc.add_paragraph()

    # 1.3 Tabla de casos de uso
    add_heading_styled(doc, "1.3 Descripción de Casos de Uso", 2)

    cu_headers = ["ID", "Caso de Uso", "Actor(es)", "Endpoint", "Postcondiciones"]
    cu_rows = [
        ["CU-01", "Registrar Log", "Microservicio Externo", "POST /api/v1/audit/log", "Log persistido, 201"],
        ["CU-02", "Registrar Batch", "Microservicio Externo", "POST /api/v1/audit/log/batch", "Logs atómicos, 201"],
        ["CU-03", "Validar API Key", "Sistema (interno)", "— (Depends)", "SHA-256 vs BD"],
        ["CU-04", "Consultar Logs", "Administrador", "GET /api/v1/audit/logs", "Lista paginada"],
        ["CU-05", "Filtrar Avanzado", "Administrador", "GET /logs?filtros", "Filtrado múltiple"],
        ["CU-06", "Obtener por ID", "Administrador", "GET /api/v1/audit/log/{id}", "Log o 404"],
        ["CU-07", "Trazar Request", "Administrador", "GET /trace/{request_id}", "Logs por X-Request-ID"],
        ["CU-08", "Logs por Usuario", "Administrador", "GET /user/{usuario_id}", "Historial paginado"],
        ["CU-09", "Buscar Full-Text", "Administrador", "GET /logs?search_text=", "Búsqueda español (GIN)"],
        ["CU-10", "Estadísticas", "Administrador", "GET /api/v1/audit/stats", "Métricas agregadas"],
        ["CU-11", "Purgar Manual", "Administrador", "DELETE /api/v1/audit/purge", "Logs eliminados"],
        ["CU-12", "Purga TTL", "Scheduler", "— (interno)", "Logs > 90d eliminados"],
        ["CU-13", "Estado Servicio", "Administrador", "GET /api/v1/audit/health", "\"is running\""],
        ["CU-14", "Health Check", "Administrador", "GET /", "Info microservicio"],
    ]
    add_table(doc, cu_headers, cu_rows)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 2. DIAGRAMA DE CLASES
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "2. Diagrama de Clases UML", 1)

    add_heading_styled(doc, "2.1 Diagrama Completo", 2)

    classes_diagram = """
╔═══════════════════════ CAPA DE PRESENTACIÓN ════════════════════════╗
║                                                                     ║
║  ┌──────────────────────────────────────────────────────────┐      ║
║  │            «controller» audit_routes                      │      ║
║  │            (APIRouter: /api/v1/audit)                     │      ║
║  ├──────────────────────────────────────────────────────────┤      ║
║  │ + router: APIRouter                                       │      ║
║  ├──────────────────────────────────────────────────────────┤      ║
║  │ + health_check() : MessageResponse                        │      ║
║  │ + create_audit_log(data, db, _api_key) : DataResponse     │      ║
║  │ + create_audit_logs_batch(logs, db, _api_key): DataResponse│     ║
║  │ + get_audit_logs(page, size, filtros, db): PaginatedResp  │      ║
║  │ + get_audit_log_by_id(audit_id, db): DataResponse         │      ║
║  │ + trace_request(request_id, db): DataResponse             │      ║
║  │ + get_user_audit_logs(usuario_id, page, db): Paginated    │      ║
║  │ + get_statistics(db): StatsResponse                       │      ║
║  │ + purge_logs(before_date, db, _api_key): MessageResponse  │      ║
║  └────────────────────────────┬─────────────────────────────┘      ║
║                               │ usa                                  ║
╚═══════════════════════════════╪═════════════════════════════════════╝
                                │
                                ▼
╔══════════════════════ CAPA DE NEGOCIO ══════════════════════════════╗
║                                                                     ║
║  ┌─────────────────────────┐  ┌───────────────────────────────┐    ║
║  │  «service» AuditService │  │ «service» StatisticsService   │    ║
║  ├─────────────────────────┤  ├───────────────────────────────┤    ║
║  │ - db: AsyncSession      │  │ - repo: AuditRepository       │    ║
║  │ - repo: AuditRepository │  ├───────────────────────────────┤    ║
║  ├─────────────────────────┤  │ + __init__(db: AsyncSession)  │    ║
║  │ + __init__(db)          │  │ + get_general_stats(): dict   │    ║
║  │ + create_log(): Response│  └───────────────────────────────┘    ║
║  │ + create_logs_batch():[]│                                        ║
║  │ + get_by_id(): Resp?    │  ┌───────────────────────────────┐    ║
║  │ + get_logs(): Paginated │  │ «service» RetentionService    │    ║
║  │ + get_by_request_id():[]│  ├───────────────────────────────┤    ║
║  │ + get_by_usuario(): Pg  │  │ - _task: asyncio.Task | None  │    ║
║  │ + purge_old_logs(): int │  │ - _running: bool              │    ║
║  └────────┬────────────────┘  ├───────────────────────────────┤    ║
║           │                    │ + start(): None               │    ║
║  ┌────────┴────────────────┐  │ + stop(): None                │    ║
║  │ «service» AuthService   │  │ + purge_old_logs(): int       │    ║
║  ├─────────────────────────┤  │ - _scheduler_loop(): None     │    ║
║  │ + validate_session()    │  │ - _seconds_until_next(): float│    ║
║  │ + check_permission()    │  └───────────────────────────────┘    ║
║  └─────────────────────────┘                                        ║
║                                                                     ║
╚═══════════════════════════════╪═════════════════════════════════════╝
                                │
                                ▼
╔══════════════════════ CAPA DE DATOS ════════════════════════════════╗
║                                                                     ║
║  ┌──────────────────────────────────────────────────────────┐      ║
║  │          «repository» AuditRepository                     │      ║
║  ├──────────────────────────────────────────────────────────┤      ║
║  │ - session: AsyncSession                                   │      ║
║  ├──────────────────────────────────────────────────────────┤      ║
║  │ + save(audit_log): AuditLog                               │      ║
║  │ + save_batch(logs): List[AuditLog]                        │      ║
║  │ + find_by_id(id): AuditLog?                               │      ║
║  │ + find_all(page, size, **filtros): (List, int)            │      ║
║  │ + find_by_request_id(rid): List[AuditLog]                 │      ║
║  │ + find_by_usuario(uid, page, size): (List, int)           │      ║
║  │ + count_total(): int                                      │      ║
║  │ + count_by_servicio(): List[tuple]                        │      ║
║  │ + count_by_codigo_respuesta(): List[tuple]                │      ║
║  │ + average_duration_by_servicio(): List[tuple]             │      ║
║  │ + error_rate_by_servicio(): List[tuple]                   │      ║
║  │ + delete_before(date): int                                │      ║
║  └──────────────────────────────────────────────────────────┘      ║
║                                                                     ║
╚═══════════════════════════════╪═════════════════════════════════════╝
                                │
                                ▼
╔═════════════════ CAPA DE DOMINIO (Modelos + Schemas) ══════════════╗
║                                                                     ║
║  ┌──────────────────────┐    ┌──────────────────────┐              ║
║  │ «entity» AuditLog    │    │ «entity» Microservice│              ║
║  │ (audit_logs)          │    │ Token                │              ║
║  ├──────────────────────┤    ├──────────────────────┤              ║
║  │ + id: UUID [PK]      │    │ + id: UUID [PK]      │              ║
║  │ + request_id: Str(50) │    │ + nombre_ms: Str(50) │              ║
║  │ + servicio: Str(50)  │    │ + token_hash: Str(256)│              ║
║  │ + endpoint: Str(200) │    │ + activo: Boolean     │              ║
║  │ + metodo: Str(10)    │    │ + created_at: TS(tz)  │              ║
║  │ + codigo_resp: Int   │    │ + updated_at: TS(tz)  │              ║
║  │ + duracion_ms: Int   │    └──────────────────────┘              ║
║  │ + usuario_id: UUID?  │                                           ║
║  │ + detalle: Text?     │    ┌──────────────────────┐              ║
║  │ + ip_origen: Str(45)?│    │ AuditLogCreate (Pyda)│              ║
║  │ + timestamp_evt: TS  │    │ AuditLogResponse     │              ║
║  │ + created_at: TS     │    │ AuditLogFilter       │              ║
║  └──────────────────────┘    │ MessageResponse      │              ║
║                               │ DataResponse<T>      │              ║
║                               │ PaginatedResponse<T> │              ║
║                               │ ErrorResponse        │              ║
║                               │ StatsResponse        │              ║
║                               └──────────────────────┘              ║
║                                                                     ║
╚═════════════════════════════════════════════════════════════════════╝
"""
    add_code_block(doc, classes_diagram)
    doc.add_paragraph()

    # 2.2 Relaciones
    add_heading_styled(doc, "2.2 Relaciones entre Clases", 2)

    rel_diagram = """
  ┌──────────────┐         ┌──────────────────┐         ┌──────────────────┐
  │ audit_routes │ ──usa──►│  AuditService    │ ──usa──►│ AuditRepository  │
  │ (Controller) │         │  (Service)       │         │ (Repository)     │
  └──────┬───────┘         └────────┬─────────┘         └────────┬─────────┘
         │                          │                             │
         │ ──usa──► StatisticsService ──usa──► AuditRepository    │
         │                                                        │
         │ ──depends──► verify_api_key() ──consulta──►            │
         │              (core/auth.py)        MicroserviceToken   │
         │                                                        │
         │ ──depends──► get_db()              │ opera sobre       │
         │              (dependencies.py)     ▼                   │
         │                           ┌──────────────────┐        │
         │                           │    AuditLog      │◄───────┘
         │                           │    (Model)       │
         │                           └──────────────────┘
         │
         │ ──valida con──►  AuditLogCreate (schema entrada)
         │ ──responde con──► AuditLogResponse (schema salida)
         │ ──filtra con──►  AuditLogFilter (schema filtros)

  RetentionService ──────► AsyncSessionLocal ──────► AuditLog (DELETE directo)
"""
    add_code_block(doc, rel_diagram)
    doc.add_paragraph()

    # 2.3 Tabla de relaciones
    add_heading_styled(doc, "2.3 Tabla de Relaciones", 2)

    rel_headers = ["Clase Origen", "Relación", "Clase Destino", "Card.", "Descripción"]
    rel_rows = [
        ["audit_routes", "usa", "AuditService", "1→1", "Cada endpoint crea instancia de servicio"],
        ["audit_routes", "usa", "StatisticsService", "1→1", "Endpoint /stats usa servicio de estadísticas"],
        ["audit_routes", "depends", "verify_api_key()", "1→1", "POST y DELETE requieren validación"],
        ["audit_routes", "depends", "get_db()", "1→1", "Inyección AsyncSession por request"],
        ["AuditService", "usa", "AuditRepository", "1→1", "Composición: servicio contiene repositorio"],
        ["StatisticsService", "usa", "AuditRepository", "1→1", "Composición: servicio contiene repositorio"],
        ["AuditRepository", "opera sobre", "AuditLog", "1→*", "CRUD sobre la entidad"],
        ["verify_api_key", "consulta", "MicroserviceToken", "1→0..1", "Busca token activo por hash"],
        ["RetentionService", "elimina", "AuditLog", "1→*", "DELETE WHERE timestamp < cutoff"],
        ["AuditLogCreate", "mapea a", "AuditLog", "1→1", "Schema → Modelo ORM"],
        ["AuditLog", "mapea a", "AuditLogResponse", "1→1", "ORM → Schema (model_validate)"],
    ]
    add_table(doc, rel_headers, rel_rows)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 3. DIAGRAMAS DE SECUENCIA
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "3. Diagramas de Secuencia UML", 1)

    # 3.1 Crear Log
    add_heading_styled(doc, "3.1 Secuencia: Crear Log de Auditoría (POST /api/v1/audit/log)", 2)

    seq1 = """
  Microservicio    RequestID MW    RateLimit MW    audit_routes    verify_api_key    AuditService    AuditRepo
  ─────┬──────     ────┬─────     ────┬──────     ────┬──────     ────┬────────     ────┬──────     ───┬─────
       │                │              │               │              │                  │             │
       │ POST /log      │              │               │              │                  │             │
       │ X-API-Key: xxx │              │               │              │                  │             │
       │ Body: {...}    │              │               │              │                  │             │
       │───────────────►│              │               │              │                  │             │
       │                │ Inyecta      │               │              │                  │             │
       │                │ X-Request-ID │               │              │                  │             │
       │                │ Inicia timer │               │              │                  │             │
       │                │─────────────►│               │              │                  │             │
       │                │              │ Verifica IP   │              │                  │             │
       │                │              │ sliding window│              │                  │             │
       │                │              │ [OK: < límite]│              │                  │             │
       │                │              │──────────────►│              │                  │             │
       │                │              │               │              │                  │             │
       │                │              │               │ Depends(     │                  │             │
       │                │              │               │ verify_api_key)                 │             │
       │                │              │               │─────────────►│                  │             │
       │                │              │               │              │ SHA-256(key)     │             │
       │                │              │               │              │ SELECT FROM      │             │
       │                │              │               │              │ microservice_tokens             │
       │                │              │               │              │ WHERE hash=? AND activo=true   │
       │                │              │               │ [API Key OK] │                  │             │
       │                │              │               │◄─────────────│                  │             │
       │                │              │               │              │                  │             │
       │                │              │               │ Pydantic valida AuditLogCreate  │             │
       │                │              │               │                                 │             │
       │                │              │               │ service.create_log(data)        │             │
       │                │              │               │────────────────────────────────►│             │
       │                │              │               │                                 │ repo.save() │
       │                │              │               │                                 │────────────►│
       │                │              │               │                                 │ session.add │
       │                │              │               │                                 │ flush()     │
       │                │              │               │                                 │ refresh()   │
       │                │              │               │                                 │◄────────────│
       │                │              │               │                                 │ db.commit() │
       │                │              │               │  AuditLogResponse               │             │
       │                │              │               │◄────────────────────────────────│             │
       │                │              │  201 Created   │                                 │             │
       │                │              │◄───────────────│                                 │             │
       │                │ + X-Request-ID               │                                 │             │
       │                │ + X-Response-Time-ms         │                                 │             │
       │                │◄──────────────               │                                 │             │
       │  201 Created    │              │               │              │                  │             │
       │  {success:true} │              │               │              │                  │             │
       │◄────────────────│              │               │              │                  │             │
"""
    add_code_block(doc, seq1)

    doc.add_page_break()

    # 3.2 Consultar Logs
    add_heading_styled(doc, "3.2 Secuencia: Consultar Logs con Filtros (GET /api/v1/audit/logs)", 2)

    seq2 = """
  Administrador      audit_routes      AuditService      AuditRepo        PostgreSQL
  ─────┬──────       ────┬──────       ────┬──────       ────┬──────      ────┬──────
       │                  │                │                  │               │
       │ GET /logs?       │                │                  │               │
       │ servicio=ms-mat  │                │                  │               │
       │ &page=1&size=20  │                │                  │               │
       │─────────────────►│                │                  │               │
       │                  │ Construye      │                  │               │
       │                  │ AuditLogFilter │                  │               │
       │                  │                │                  │               │
       │                  │ service.get_logs()                │               │
       │                  │───────────────►│                  │               │
       │                  │                │ repo.find_all()  │               │
       │                  │                │ (page, filters)  │               │
       │                  │                │─────────────────►│               │
       │                  │                │                  │ COUNT(id)     │
       │                  │                │                  │ WHERE srv=?   │
       │                  │                │                  │──────────────►│
       │                  │                │                  │◄── total:150 ─│
       │                  │                │                  │               │
       │                  │                │                  │ SELECT * FROM │
       │                  │                │                  │ audit_logs    │
       │                  │                │                  │ WHERE srv=?   │
       │                  │                │                  │ ORDER BY ts   │
       │                  │                │                  │ OFFSET 0      │
       │                  │                │                  │ LIMIT 20      │
       │                  │                │                  │──────────────►│
       │                  │                │                  │◄── [20 rows] ─│
       │                  │                │◄── (results,150) │               │
       │                  │                │ model_validate() │               │
       │                  │  Paginated     │ por cada result  │               │
       │                  │  {data, total: │                  │               │
       │                  │   150, page:1, │                  │               │
       │                  │   pages:8}     │                  │               │
       │                  │◄───────────────│                  │               │
       │  200 OK          │                │                  │               │
       │  {success:true,  │                │                  │               │
       │   data:[...],    │                │                  │               │
       │   total_pages:8} │                │                  │               │
       │◄─────────────────│                │                  │               │
"""
    add_code_block(doc, seq2)

    doc.add_paragraph()

    # 3.3 Scheduler Retention
    add_heading_styled(doc, "3.3 Secuencia: Scheduler de Retención Automática (TTL)", 2)

    seq3 = """
  FastAPI Lifespan    RetentionService    AsyncSessionLocal    PostgreSQL
  ─────┬──────        ─────┬──────        ─────┬──────         ────┬──────
       │                    │                   │                    │
       │ ══ STARTUP ══      │                   │                    │
       │                    │                   │                    │
       │  start()           │                   │                    │
       │───────────────────►│                   │                    │
       │                    │ _running = True   │                    │
       │                    │ asyncio.create_   │                    │
       │                    │ task(scheduler)   │                    │
       │  ◄── return ───────│                   │                    │
       │                    │                   │                    │
       │   ══ SCHEDULER LOOP (background) ══   │                    │
       │                    │                   │                    │
       │                    │ _seconds_until_   │                    │
       │                    │ next_run()        │                    │
       │                    │ (ej: 28800 seg)   │                    │
       │                    │                   │                    │
       │                    │ asyncio.sleep(    │                    │
       │                    │   28800)          │                    │
       │                    │ ... espera hasta  │                    │
       │                    │ 03:00 UTC ...     │                    │
       │                    │                   │                    │
       │                    │ ══ 03:00 UTC ══   │                    │
       │                    │ purge_old_logs()  │                    │
       │                    │                   │                    │
       │                    │ cutoff = now() -  │                    │
       │                    │ timedelta(90 days)│                    │
       │                    │                   │                    │
       │                    │ Abre sesión propia│                    │
       │                    │──────────────────►│                    │
       │                    │                   │ DELETE FROM        │
       │                    │                   │ audit_logs WHERE   │
       │                    │                   │ timestamp_evento   │
       │                    │                   │ < cutoff_date      │
       │                    │                   │───────────────────►│
       │                    │                   │◄── rowcount:1523 ──│
       │                    │                   │ COMMIT             │
       │                    │                   │───────────────────►│
       │                    │                   │◄── OK ─────────────│
       │                    │◄── 1523 eliminados│                    │
       │                    │ logger.info(      │                    │
       │                    │  "purge_completed"│                    │
       │                    │  deleted: 1523)   │                    │
       │                    │                   │                    │
       │                    │ (vuelve al loop)  │                    │
       │                    │                   │                    │
       │  ══ SHUTDOWN ══    │                   │                    │
       │  stop()            │                   │                    │
       │───────────────────►│                   │                    │
       │                    │ _running = False  │                    │
       │                    │ task.cancel()     │                    │
"""
    add_code_block(doc, seq3)

    doc.add_page_break()

    # 3.4 Validación API Key
    add_heading_styled(doc, "3.4 Secuencia: Validación de API Key (Detalle)", 2)

    seq4 = """
  Request (Depends)    verify_api_key     AsyncSession      PostgreSQL
  ─────┬──────         ─────┬──────       ─────┬──────      ────┬──────
       │                     │                  │                │
       │ Header: X-API-Key   │                  │                │
       │────────────────────►│                  │                │
       │                     │ api_key == None?  │                │
       │                     │                  │                │
       │  ALT [No Key + dev/testing]            │                │
       │  ┌─────────────────────────────────────────────────────┐
       │  │                  │ return None       │                │
       │  │◄─────────────────│ (acceso permitido)│                │
       │  └─────────────────────────────────────────────────────┘
       │                     │                  │                │
       │  ALT [No Key + production]             │                │
       │  ┌─────────────────────────────────────────────────────┐
       │  │                  │ HTTPException     │                │
       │  │◄─────────────────│ (401 Unauthorized)│                │
       │  └─────────────────────────────────────────────────────┘
       │                     │                  │                │
       │  ALT [API Key presente]                │                │
       │  ┌─────────────────────────────────────────────────────┐
       │  │                  │ hash = SHA-256(key)               │
       │  │                  │                  │                │
       │  │                  │ SELECT * FROM    │                │
       │  │                  │ microservice_    │                │
       │  │                  │ tokens WHERE     │                │
       │  │                  │ hash=? AND       │                │
       │  │                  │ activo=True      │                │
       │  │                  │─────────────────►│               │
       │  │                  │                  │──────────────►│
       │  │                  │                  │◄──────────────│
       │  │                  │◄─────────────────│               │
       │  │                  │                  │                │
       │  │  [found]         │ return token     │                │
       │  │◄─────────────────│                  │                │
       │  │  [not found]     │ HTTPException 401│                │
       │  │◄─────────────────│                  │                │
       │  └─────────────────────────────────────────────────────┘
"""
    add_code_block(doc, seq4)

    doc.add_paragraph()

    # 3.5 Estadísticas
    add_heading_styled(doc, "3.5 Secuencia: Obtener Estadísticas (GET /api/v1/audit/stats)", 2)

    seq5 = """
  Administrador    audit_routes    StatisticsService    AuditRepo       PostgreSQL
  ─────┬──────     ────┬──────     ─────┬──────         ────┬──────     ────┬──────
       │                │               │                    │              │
       │ GET /stats     │               │                    │              │
       │───────────────►│               │                    │              │
       │                │ get_general_  │                    │              │
       │                │ stats()       │                    │              │
       │                │──────────────►│                    │              │
       │                │               │ count_total()      │              │
       │                │               │───────────────────►│── COUNT(*) ─►│
       │                │               │◄── total: 5000 ────│◄─────────────│
       │                │               │                    │              │
       │                │               │ count_by_servicio()│              │
       │                │               │───────────────────►│─ GROUP BY ──►│
       │                │               │◄── [{srv, cnt}] ───│◄─────────────│
       │                │               │                    │              │
       │                │               │ count_by_codigo()  │              │
       │                │               │───────────────────►│─ GROUP BY ──►│
       │                │               │◄── [{cod, cnt}] ───│◄─────────────│
       │                │               │                    │              │
       │                │               │ avg_duration()     │              │
       │                │               │───────────────────►│─ AVG(dur) ──►│
       │                │               │◄── [{srv, avg}] ───│◄─────────────│
       │                │               │                    │              │
       │                │               │ error_rate()       │              │
       │                │               │───────────────────►│─ SUM/CASE ──►│
       │                │               │◄── [{srv,rate}] ───│◄─────────────│
       │                │               │                    │              │
       │                │  StatsResponse│                    │              │
       │                │  {total, srv, │                    │              │
       │                │   cod, avg,   │                    │              │
       │                │   errores}    │                    │              │
       │                │◄──────────────│                    │              │
       │  200 OK        │               │                    │              │
       │  {success:true}│               │                    │              │
       │◄───────────────│               │                    │              │
"""
    add_code_block(doc, seq5)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 4. DIAGRAMA DE COMPONENTES
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "4. Diagrama de Componentes Interno", 1)

    add_heading_styled(doc, "4.1 Diagrama de Componentes (Clean Architecture)", 2)

    comp_diagram = """
╔═════════════════════════════════════════════════════════════════════════════╗
║                        ms-auditoria (:8019)                               ║
║                                                                           ║
║  ┌─────────────── «component» MIDDLEWARE CHAIN ───────────────────┐       ║
║  │  [RequestID MW] ──► [RateLimit MW] ──► [CORS MW]              │       ║
║  │  • X-Request-ID     • 100 req/60s      • Orígenes             │       ║
║  │  • Response time    • Sliding window   • Methods              │       ║
║  └────────────────────────────────┬───────────────────────────────┘       ║
║                                   │                                       ║
║  ┌──────── «component» EXCEPTION HANDLERS ────────────────────────┐       ║
║  │  [HTTP 4xx/5xx]  [Validation 422]  [Unhandled 500]            │       ║
║  └────────────────────────────────┬───────────────────────────────┘       ║
║                                   │                                       ║
║                                   ▼                                       ║
║  ┌─────────────── «component» PRESENTACIÓN ───────────────────────┐       ║
║  │  audit_routes.py (APIRouter /api/v1/audit/*)                   │       ║
║  │  9 endpoints + Depends(verify_api_key, get_db)                 │       ║
║  └────────────────────────────────┬───────────────────────────────┘       ║
║                                   │                                       ║
║                                   ▼                                       ║
║  ┌─────────────── «component» LÓGICA DE NEGOCIO ─────────────────┐       ║
║  │  [AuditService]  [StatisticsService]  [RetentionService]      │       ║
║  │  [AuthService]   [AESCipher (disponible)]                     │       ║
║  └────────────────────────────────┬───────────────────────────────┘       ║
║                                   │                                       ║
║                                   ▼                                       ║
║  ┌─────────────── «component» ACCESO A DATOS ────────────────────┐       ║
║  │  AuditRepository (12 métodos: CRUD + stats + delete)          │       ║
║  └────────────────────────────────┬───────────────────────────────┘       ║
║                                   │                                       ║
║                                   ▼                                       ║
║  ┌─────────────── «component» BASE DE DATOS ─────────────────────┐       ║
║  │  [connection.py: async_engine + pool]  [session.py: AsyncSess] │       ║
║  │  [base.py: DeclarativeBase]  [config.py: Settings + env vars] │       ║
║  └────────────────────────────────┬───────────────────────────────┘       ║
║                                   │                                       ║
║  ┌─────── «component» DOMINIO ────┼────────────── «component» UTILS ─┐   ║
║  │  AuditLog, MicroserviceToken,  │   logger.py (JSON Structured)    │   ║
║  │  GUID, Schemas (Pydantic)      │   JSONFormatter → stdout         │   ║
║  └────────────────────────────────┼──────────────────────────────────┘   ║
║                                   │                                       ║
╚═══════════════════════════════════╪═══════════════════════════════════════╝
                                    │ asyncpg (pool_size=10, max_overflow=20)
                                    ▼
                        ┌──────────────────────┐
                        │    PostgreSQL 16      │
                        │    (:5432)            │
                        │  • audit_logs         │
                        │  • microservice_tokens│
                        │  • 8 índices + GIN    │
                        └──────────────────────┘
"""
    add_code_block(doc, comp_diagram)
    doc.add_paragraph()

    # 4.2 Simplificado
    add_heading_styled(doc, "4.2 Diagrama Simplificado de Capas", 2)

    simple_diagram = """
                        ┌───────────────────────┐
                        │       REQUEST         │
                        └───────────┬───────────┘
                                    │
                                    ▼
                        ┌───────────────────────┐
                        │      MIDDLEWARE       │
                        │  RequestID → RateLimit │
                        │  → CORS               │
                        └───────────┬───────────┘
                                    │
                                    ▼
                        ┌───────────────────────┐
                        │       ROUTER         │
                        │   (audit_routes.py)   │
                        │   9 endpoints         │
                        │   + Depends(auth, db) │
                        └───────────┬───────────┘
                                    │
                                    ▼
                        ┌───────────────────────┐
                        │      SERVICE         │
                        │  audit_service.py     │
                        │  statistics_service   │
                        │  retention_service    │
                        └───────────┬───────────┘
                                    │
                                    ▼
                        ┌───────────────────────┐
                        │     REPOSITORY       │
                        │ audit_repository.py   │
                        │ SQLAlchemy 2.0 async  │
                        └───────────┬───────────┘
                                    │
                                    ▼
                        ┌───────────────────────┐
                        │      DATABASE        │
                        │  AsyncSession +       │
                        │  AsyncEngine (asyncpg)│
                        └───────────┬───────────┘
                                    │
                                    ▼
                        ┌───────────────────────┐
                        │    PostgreSQL 16      │
                        │  audit_logs           │
                        │  microservice_tokens  │
                        └───────────────────────┘
"""
    add_code_block(doc, simple_diagram)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 5. DIAGRAMA ENTIDAD-RELACIÓN
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "5. Diagrama Entidad-Relación (ER)", 1)

    add_heading_styled(doc, "5.1 Diagrama ER Completo", 2)

    er_diagram = """
╔═══════════════════════════════════════════════════════════════════════════════╗
║                      DIAGRAMA ENTIDAD-RELACIÓN                              ║
║                      ms-auditoria — PostgreSQL 16                           ║
║                                                                             ║
║  ┌───────────────────────────────────────────────────────┐                 ║
║  │                     audit_logs                         │                 ║
║  │                     (Tabla principal)                  │                 ║
║  ├───────────────────────────────────────────────────────┤                 ║
║  │ «PK» id               UUID            NOT NULL        │                 ║
║  │───────────────────────────────────────────────────────│                 ║
║  │      request_id       VARCHAR(50)     NOT NULL        │ ◄── IX         ║
║  │      servicio         VARCHAR(50)     NOT NULL        │ ◄── IX         ║
║  │      endpoint         VARCHAR(200)    NOT NULL        │                 ║
║  │      metodo           VARCHAR(10)     NOT NULL        │                 ║
║  │      codigo_respuesta INTEGER         NULLABLE        │ ◄── IX         ║
║  │      duracion_ms      INTEGER         NULLABLE        │                 ║
║  │      usuario_id       UUID            NULLABLE        │ ◄── IX         ║
║  │      detalle          TEXT            NULLABLE        │ ◄── GIN(FTS)   ║
║  │      ip_origen        VARCHAR(45)     NULLABLE        │                 ║
║  │      timestamp_evento TIMESTAMP(tz)   NOT NULL        │                 ║
║  │      created_at       TIMESTAMP(tz)   NOT NULL        │                 ║
║  ├───────────────────────────────────────────────────────┤                 ║
║  │  Índices compuestos:                                  │                 ║
║  │  • (servicio, timestamp_evento)                       │                 ║
║  │  • (usuario_id, timestamp_evento)                     │                 ║
║  │  • (codigo_respuesta, servicio)                       │                 ║
║  │  Índice GIN: to_tsvector('spanish', detalle)          │                 ║
║  └───────────────────────────────────────────────────────┘                 ║
║                                                                             ║
║         Sin FK directa — relación lógica vía API Key en runtime            ║
║                                                                             ║
║  ┌───────────────────────────────────────────────────────┐                 ║
║  │                  microservice_tokens                   │                 ║
║  │                  (Tabla de autenticación)              │                 ║
║  ├───────────────────────────────────────────────────────┤                 ║
║  │ «PK» id                     UUID          NOT NULL    │                 ║
║  │───────────────────────────────────────────────────────│                 ║
║  │ «UQ» nombre_microservicio   VARCHAR(50)   NOT NULL    │                 ║
║  │      token_hash             VARCHAR(256)  NOT NULL    │                 ║
║  │      activo                 BOOLEAN       NOT NULL    │                 ║
║  │      created_at             TIMESTAMP(tz) NOT NULL    │                 ║
║  │      updated_at             TIMESTAMP(tz) NOT NULL    │                 ║
║  └───────────────────────────────────────────────────────┘                 ║
║                                                                             ║
║  Relación lógica: microservice_tokens 1 ─────── 0..* audit_logs           ║
║  (nombre_microservicio ≈ servicio, validada en runtime, no FK en BD)       ║
║                                                                             ║
╚═══════════════════════════════════════════════════════════════════════════════╝
"""
    add_code_block(doc, er_diagram)
    doc.add_paragraph()

    # 5.2 ER con cardinalidades
    add_heading_styled(doc, "5.2 Notación ER con Cardinalidades", 2)

    er_card = """
  ┌───────────────────┐          1 : 0..*          ┌───────────────────┐
  │                   │                             │                   │
  │ microservice_     │────────────────────────────►│   audit_logs      │
  │ tokens            │  "Un microservicio          │                   │
  │                   │   registrado genera         │                   │
  │ PK: id (UUID)     │   0 o muchos logs"          │ PK: id (UUID)     │
  │ UQ: nombre_ms     │                             │    request_id     │
  │    token_hash     │   Relación lógica vía       │    servicio       │
  │    activo         │   API Key en runtime        │    endpoint       │
  │    created_at     │   (NO es FK en BD)          │    metodo         │
  │    updated_at     │                             │    codigo_resp    │
  │                   │                             │    duracion_ms    │
  └───────────────────┘                             │    usuario_id     │
                                                    │    detalle        │
                                                    │    ip_origen      │
                                                    │    timestamp_evt  │
                                                    │    created_at     │
                                                    │                   │
                                                    └───────────────────┘
"""
    add_code_block(doc, er_card)
    doc.add_paragraph()

    # 5.3 Tablas de atributos
    add_heading_styled(doc, "5.3 Detalle de Atributos con Tipos y Restricciones", 2)

    doc.add_paragraph().add_run("Tabla: audit_logs").bold = True
    doc.add_paragraph()

    al_headers = ["#", "Atributo", "Tipo PostgreSQL", "PK", "NN", "Default", "Descripción"]
    al_rows = [
        ["1", "id", "UUID", "✅", "✅", "uuid4()", "Identificador único"],
        ["2", "request_id", "VARCHAR(50)", "", "✅", "—", "X-Request-ID trazabilidad"],
        ["3", "servicio", "VARCHAR(50)", "", "✅", "—", "Microservicio emisor"],
        ["4", "endpoint", "VARCHAR(200)", "", "✅", "—", "Ruta invocada"],
        ["5", "metodo", "VARCHAR(10)", "", "✅", "—", "HTTP method"],
        ["6", "codigo_respuesta", "INTEGER", "", "", "—", "Código HTTP"],
        ["7", "duracion_ms", "INTEGER", "", "", "—", "Duración en ms"],
        ["8", "usuario_id", "UUID", "", "", "None", "UUID usuario"],
        ["9", "detalle", "TEXT", "", "", "None", "Detalle JSON string"],
        ["10", "ip_origen", "VARCHAR(45)", "", "", "None", "IP origen (v4/v6)"],
        ["11", "timestamp_evento", "TIMESTAMP(tz)", "", "✅", "—", "Momento del evento"],
        ["12", "created_at", "TIMESTAMP(tz)", "", "✅", "now(UTC)", "Registro en BD"],
    ]
    add_table(doc, al_headers, al_rows)
    doc.add_paragraph()

    doc.add_paragraph().add_run("Tabla: microservice_tokens").bold = True
    doc.add_paragraph()

    mt_headers = ["#", "Atributo", "Tipo PostgreSQL", "PK", "UQ", "NN", "Default", "Descripción"]
    mt_rows = [
        ["1", "id", "UUID", "✅", "", "✅", "uuid4()", "Identificador único"],
        ["2", "nombre_microservicio", "VARCHAR(50)", "", "✅", "✅", "—", "Nombre del ms"],
        ["3", "token_hash", "VARCHAR(256)", "", "", "✅", "—", "SHA-256 del API Key"],
        ["4", "activo", "BOOLEAN", "", "", "✅", "True", "Autorizado o no"],
        ["5", "created_at", "TIMESTAMP(tz)", "", "", "✅", "now(UTC)", "Creación"],
        ["6", "updated_at", "TIMESTAMP(tz)", "", "", "✅", "now(UTC)", "Última actualización"],
    ]
    add_table(doc, mt_headers, mt_rows)
    doc.add_paragraph()

    # 5.4 Resumen de Índices
    add_heading_styled(doc, "5.4 Resumen de Índices", 2)

    idx_headers = ["#", "Nombre", "Tabla", "Tipo", "Columna(s)", "Propósito"]
    idx_rows = [
        ["1", "PK (id)", "audit_logs", "B-Tree (PK)", "id", "Clave primaria"],
        ["2", "ix_audit_logs_request_id", "audit_logs", "B-Tree", "request_id", "Trazabilidad"],
        ["3", "ix_audit_logs_servicio", "audit_logs", "B-Tree", "servicio", "Filtro por servicio"],
        ["4", "ix_audit_logs_codigo_resp", "audit_logs", "B-Tree", "codigo_respuesta", "Filtro por código"],
        ["5", "ix_audit_logs_usuario_id", "audit_logs", "B-Tree", "usuario_id", "Filtro por usuario"],
        ["6", "ix_audit_srv_timestamp", "audit_logs", "Compuesto", "servicio, timestamp", "Rango temporal"],
        ["7", "ix_audit_usr_timestamp", "audit_logs", "Compuesto", "usuario_id, timestamp", "Rango temporal"],
        ["8", "ix_audit_cod_servicio", "audit_logs", "Compuesto", "codigo_resp, servicio", "Estadísticas"],
        ["9", "ix_audit_detalle_fts", "audit_logs", "GIN", "tsvector(detalle)", "Full-text español"],
        ["10", "PK (id)", "ms_tokens", "B-Tree (PK)", "id", "Clave primaria"],
        ["11", "UQ (nombre)", "ms_tokens", "B-Tree (UQ)", "nombre_microservicio", "Unicidad"],
    ]
    add_table(doc, idx_headers, idx_rows)

    # ══════════════════════════════════════════════════════════════
    # GUARDAR
    # ══════════════════════════════════════════════════════════════
    output_path = os.path.join(os.path.dirname(__file__), "Diagramas_UML_ms_auditoria.docx")
    doc.save(output_path)
    print(f"✅ Documento generado: {output_path}")
    print(f"   Secciones: 5 diagramas UML")
    print(f"   Formato: Word (.docx), Landscape A4")


if __name__ == "__main__":
    generate_docx()
