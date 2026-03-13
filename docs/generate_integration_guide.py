"""
Generador de Guía de Integración para ms-auditoria
Produce: ms-auditoria_Guia_Integracion.docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "ms-auditoria_Guia_Integracion.docx")

# ─── COLORES ────────────────────────────────────────────────────────────────
C_TITLE       = RGBColor(0x1A, 0x52, 0x76)   # Azul oscuro corporativo
C_H1          = RGBColor(0x1A, 0x52, 0x76)
C_H2          = RGBColor(0x21, 0x6A, 0x94)
C_H3          = RGBColor(0x2E, 0x86, 0xC1)
C_CODE_BG     = RGBColor(0xF0, 0xF4, 0xF8)   # Gris muy claro
C_CODE_BORDER = RGBColor(0x5D, 0xAD, 0xE8)   # Azul claro borde
C_TABLE_HEAD  = RGBColor(0x1A, 0x52, 0x76)   # Cabecera tabla
C_TABLE_ALT   = RGBColor(0xEB, 0xF5, 0xFB)   # Filas alternas
C_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
C_WARN_BG     = RGBColor(0xFF, 0xF9, 0xE6)
C_OK_BG       = RGBColor(0xE9, 0xF7, 0xEF)
C_BODY        = RGBColor(0x1C, 0x1C, 0x1E)

# ─── HELPERS ────────────────────────────────────────────────────────────────

def rgb_to_hex(rgb: RGBColor) -> str:
    return '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = rgb_to_hex(rgb)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_paragraph_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def add_bookmark(para, name):
    """Agrega un marcador de Word a un párrafo (para hipervínculos internos)."""
    run = para.add_run()
    tag_start = OxmlElement('w:bookmarkStart')
    tag_start.set(qn('w:id'), str(abs(hash(name)) % 10000))
    tag_start.set(qn('w:name'), name)
    run._r.append(tag_start)
    tag_end = OxmlElement('w:bookmarkEnd')
    tag_end.set(qn('w:id'), str(abs(hash(name)) % 10000))
    run._r.append(tag_end)


def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    set_paragraph_spacing(para, 0, 0)


def add_horizontal_rule(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E86C1')
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_paragraph_spacing(para, 0, 60)


# ─── ESTILOS DE TEXTO ────────────────────────────────────────────────────────

def add_cover_title(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = C_TITLE
    set_paragraph_spacing(para, 0, 80)
    return para


def add_cover_subtitle(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = False
    run.font.size = Pt(16)
    run.font.color.rgb = C_H2
    set_paragraph_spacing(para, 0, 40)
    return para


def add_cover_info(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    set_paragraph_spacing(para, 0, 20)
    return para


def heading1(doc, text, bookmark=None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = C_H1
    set_paragraph_spacing(para, 240, 80)
    if bookmark:
        add_bookmark(para, bookmark)
    add_horizontal_rule(doc)
    return para


def heading2(doc, text, bookmark=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = C_H2
    set_paragraph_spacing(para, 180, 60)
    if bookmark:
        add_bookmark(para, bookmark)
    return para


def heading3(doc, text, bookmark=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = C_H3
    set_paragraph_spacing(para, 120, 40)
    if bookmark:
        add_bookmark(para, bookmark)
    return para


def body_text(doc, text, bold=False, italic=False, size=11, color=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color if color else C_BODY
    set_paragraph_spacing(para, 0, 80)
    return para


def body_text_inline(doc, parts, size=11):
    """parts = list of (text, bold, italic, color)"""
    para = doc.add_paragraph()
    for text, bold, italic, color in parts:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = color if color else C_BODY
    set_paragraph_spacing(para, 0, 80)
    return para


def bullet(doc, text, level=0, bold=False):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.color.rgb = C_BODY
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(360 + level * 360))
    pPr.append(ind)
    set_paragraph_spacing(para, 0, 40)
    return para


def numbered(doc, text, bold=False):
    para = doc.add_paragraph(style='List Number')
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.color.rgb = C_BODY
    set_paragraph_spacing(para, 0, 40)
    return para


def callout(doc, text, kind='info'):
    """kind: info, warning, success, error"""
    colors = {
        'info':    (RGBColor(0xD6, 0xEA, 0xF8), RGBColor(0x21, 0x6A, 0x94), '💡'),
        'warning': (C_WARN_BG,                   RGBColor(0x9A, 0x6C, 0x00), '⚠️'),
        'success': (C_OK_BG,                     RGBColor(0x1E, 0x8B, 0x4C), '✅'),
        'error':   (RGBColor(0xFD, 0xED, 0xEC), RGBColor(0xAB, 0x2D, 0x2D), '❌'),
    }
    bg, text_color, icon = colors.get(kind, colors['info'])
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    border_color = rgb_to_hex(text_color)
    for side in ['top', 'bottom', 'left', 'right']:
        set_cell_border(cell, **{side: {'val': 'single', 'sz': 4, 'color': border_color}})
    para = cell.paragraphs[0]
    run = para.add_run(f'{icon}  {text}')
    run.font.size = Pt(11)
    run.font.color.rgb = text_color
    set_paragraph_spacing(para, 60, 60)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 40)
    return table


def code_block(doc, code_text, lang_label=None):
    """Bloque de código con fondo azul muy claro, fuente monoespaciada."""
    outer = doc.add_table(rows=1, cols=1)
    outer.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = outer.cell(0, 0)
    set_cell_bg(cell, C_CODE_BG)
    for side in ['top', 'bottom', 'left', 'right']:
        set_cell_border(cell, **{side: {'val': 'single', 'sz': 4, 'color': '5DADE8'}})

    if lang_label:
        label_para = cell.add_paragraph()
        lr = label_para.add_run(f'  {lang_label}')
        lr.font.size = Pt(8)
        lr.font.color.rgb = RGBColor(0x5D, 0xAD, 0xE8)
        lr.bold = True
        set_paragraph_spacing(label_para, 40, 0)

    # Limpiar primer párrafo vacío
    first_para = cell.paragraphs[0]
    if not first_para.text and not lang_label:
        pass  # lo usamos como contenedor

    lines = code_text.strip().split('\n')
    for i, line in enumerate(lines):
        if i == 0 and not lang_label:
            para = first_para
        else:
            para = cell.add_paragraph()
        run = para.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x0D, 0x47, 0x6B)
        top = 20 if i > 0 else 60
        bot = 20 if i < len(lines) - 1 else 60
        set_paragraph_spacing(para, top, bot)

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, 0, 60)
    return outer


def info_table(doc, rows_data, headers=None, col_widths=None):
    """Tabla de información con cabecera azul y filas alternas."""
    n_cols = len(headers) if headers else len(rows_data[0])
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    if headers:
        row = table.add_row()
        for j, h in enumerate(headers):
            cell = row.cells[j]
            set_cell_bg(cell, C_TABLE_HEAD)
            para = cell.paragraphs[0]
            run = para.add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = C_WHITE
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(para, 40, 40)

    for i, row_data in enumerate(rows_data):
        row = table.add_row()
        bg = C_TABLE_ALT if i % 2 == 1 else C_WHITE
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            # Soporte para texto enriquecido: si es lista de (text, bold, mono)
            if isinstance(cell_text, list):
                for txt, bold, mono in cell_text:
                    run = para.add_run(txt)
                    run.bold = bold
                    if mono:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    else:
                        run.font.size = Pt(10)
                    run.font.color.rgb = C_BODY
            else:
                run = para.add_run(str(cell_text))
                run.font.size = Pt(10)
                run.font.color.rgb = C_BODY
            set_paragraph_spacing(para, 30, 30)

    if col_widths:
        for row in table.rows:
            for j, width in enumerate(col_widths):
                if j < len(row.cells):
                    row.cells[j].width = Inches(width)

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, 0, 60)
    return table


# ════════════════════════════════════════════════════════════════════════════
#  CONTENIDO DEL DOCUMENTO
# ════════════════════════════════════════════════════════════════════════════

def build_cover(doc):
    # Espaciado superior
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, 0, 0)

    add_cover_title(doc, "Guía de Integración")
    add_cover_subtitle(doc, "ms-auditoria [AUD]")
    add_cover_subtitle(doc, "ERP Universitario — Universidad del Valle, Sede Caicedonia")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, 80, 80)

    add_cover_info(doc, "Versión 1.0  ·  Marzo 2026")
    add_cover_info(doc, "FastAPI + Python + PostgreSQL")
    add_cover_info(doc, "Para uso interno del equipo de desarrollo")
    add_cover_info(doc, "Destinatarios: desarrolladores de los 18 microservicios del sistema")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, 60, 60)

    add_horizontal_rule(doc)

    callout(doc,
        "Este documento es la referencia oficial para integrar cualquier microservicio del "
        "ERP Universitario con ms-auditoria. Contiene la especificación completa del contrato "
        "de comunicación, ejemplos de código para Python/FastAPI y una guía paso a paso por "
        "cada uno de los 18 microservicios del sistema.",
        kind='info')

    add_page_break(doc)


def build_toc(doc):
    heading1(doc, "Tabla de Contenido")
    toc_items = [
        ("1.", "Introducción y Propósito de este Documento"),
        ("2.", "¿Qué es ms-auditoria y para qué sirve?"),
        ("3.", "Principios Fundamentales de la Integración"),
        ("4.", "Información de Conexión y Endpoints"),
        ("5.", "Autenticación: El Token de Aplicación"),
        ("6.", "Estructura del Registro de Log (Contrato de Datos)"),
        ("7.", "Endpoint Principal: POST /api/v1/logs"),
        ("8.", "Endpoint para Lotes: POST /api/v1/logs/batch"),
        ("9.", "Manejo de Errores y Respuestas"),
        ("10.", "Patrón de Implementación Recomendado (Python / FastAPI)"),
        ("11.", "Request ID: Trazabilidad Distribuida"),
        ("12.", "Health Check"),
        ("13.", "Ejemplos de Integración por Microservicio (18 servicios)"),
        ("14.", "Guía de Resolución de Problemas"),
        ("15.", "Resumen de Reglas y Checklist de Integración"),
    ]
    for num, title in toc_items:
        para = doc.add_paragraph()
        r1 = para.add_run(f"  {num}  ")
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = C_H2
        r2 = para.add_run(title)
        r2.font.size = Pt(11)
        r2.font.color.rgb = C_BODY
        set_paragraph_spacing(para, 20, 20)
    add_page_break(doc)


def build_sec1_intro(doc):
    heading1(doc, "1. Introducción y Propósito de este Documento")

    body_text(doc,
        "Este documento es la referencia técnica oficial para que cualquier desarrollador del "
        "ERP Universitario pueda integrar su microservicio con ms-auditoria de forma correcta, "
        "eficiente y sin afectar el rendimiento de su propio servicio.")

    body_text(doc,
        "El ERP Universitario está compuesto por 19 microservicios independientes. Todos ellos, "
        "sin excepción, tienen la obligación de registrar sus operaciones en ms-auditoria. Esta "
        "obligación está establecida en la Regla Transversal 6.6 del documento de arquitectura "
        "del sistema:", italic=True)

    callout(doc,
        '"Cada operación realizada en cualquier microservicio debe generar un registro de log '
        'en formato JSON que contenga: la fecha y hora de la operación, el identificador de '
        'rastreo de la petición, el nombre del microservicio, la funcionalidad ejecutada, el '
        'método utilizado, el código de respuesta, la duración en milisegundos, el identificador '
        'del usuario que realizó la operación y un detalle descriptivo. Estos registros se envían '
        'de forma asíncrona al servicio de auditoría, de manera que el envío no bloquee ni retrase '
        'la respuesta al usuario. Si el envío al servicio de auditoría falla, el microservicio debe '
        'continuar operando normalmente."', kind='warning')

    body_text(doc, "Este documento cubre:")
    bullet(doc, "Qué es ms-auditoria y cuál es su rol en el sistema.")
    bullet(doc, "Los principios obligatorios de integración.")
    bullet(doc, "La especificación completa del contrato de datos (campos del log, tipos, validaciones).")
    bullet(doc, "Los dos endpoints disponibles para enviar logs: individual y por lote.")
    bullet(doc, "Cómo implementar la integración en Python/FastAPI con código de ejemplo listo para usar.")
    bullet(doc, "Ejemplos concretos de logs para cada uno de los 18 microservicios del sistema.")
    bullet(doc, "Una guía de resolución de problemas frecuentes.")
    bullet(doc, "Un checklist final de verificación.")

    callout(doc,
        "Si ya implementaste la integración, usa la Sección 15 como checklist de verificación "
        "rápida antes de hacer el merge a main.",
        kind='success')


def build_sec2_que_es(doc):
    add_page_break(doc)
    heading1(doc, "2. ¿Qué es ms-auditoria y para qué sirve?")

    body_text(doc,
        "ms-auditoria es el servicio centralizado de trazabilidad y auditoría del ERP Universitario. "
        "Su propósito es almacenar un registro permanente de todas las operaciones realizadas en el "
        "sistema, permitir reconstruir la cadena completa de eventos de cualquier petición distribuida "
        "y generar estadísticas de uso por microservicio.")

    heading2(doc, "2.1 Rol en la arquitectura")
    info_table(doc,
        [
            ["Módulo", "Módulo 6 — Transversales"],
            ["Código", "AUD"],
            ["Stack", "FastAPI + Python + PostgreSQL"],
            ["Tipo de servicio", "Receptor pasivo de logs + Motor de trazabilidad"],
            ["Consumidores", "Los 18 microservicios restantes del ERP (todos)"],
            ["Dependencias externas", "ms-autenticacion (para consultas de usuarios), ms-roles (para consultas de usuarios)"],
            ["Dependencias para envío de logs", "Ninguna — el envío de logs NO requiere autenticación de usuario"],
        ],
        headers=["Campo", "Valor"],
        col_widths=[2.2, 4.0],
    )

    heading2(doc, "2.2 Lo que ms-auditoria hace por el sistema")
    bullet(doc, "Recibe logs de todos los microservicios de forma asíncrona (fire-and-forget).")
    bullet(doc, "Almacena cada registro en PostgreSQL con todos los campos estándar.")
    bullet(doc, "Permite consultar la traza completa de cualquier petición por su Request ID.")
    bullet(doc, "Permite filtrar registros por microservicio y rango de fechas.")
    bullet(doc, "Calcula estadísticas de uso: total de peticiones, errores, tiempo promedio.")
    bullet(doc, "Se auto-audita: sus propias operaciones también generan registros.")
    bullet(doc, "Rota automáticamente los registros antiguos (por defecto, mayores a 30 días).")

    heading2(doc, "2.3 Lo que ms-auditoria NO hace")
    callout(doc,
        "ms-auditoria NO valida la sesión del microservicio emisor. Para enviar logs, "
        "solo necesitas un token de aplicación (X-App-Token). NO necesitas un token de sesión "
        "de usuario (Authorization: Bearer). La validación de sesión solo aplica para "
        "operaciones de consulta realizadas por usuarios humanos.",
        kind='warning')
    bullet(doc, "NO bloquea al microservicio emisor: siempre responde de inmediato con 202 Accepted.")
    bullet(doc, "NO debe convertirse en un punto de fallo: si ms-auditoria falla, tu servicio debe seguir funcionando.")
    bullet(doc, "NO almacena credenciales, contraseñas ni tokens de sesión en los logs.")


def build_sec3_principios(doc):
    add_page_break(doc)
    heading1(doc, "3. Principios Fundamentales de la Integración")

    callout(doc,
        "Estos principios son OBLIGATORIOS para todos los microservicios del sistema. "
        "No son opcionales ni sugerencias de buenas prácticas.",
        kind='error')

    heading2(doc, "Principio 1 — El envío es fire-and-forget (asíncrono)")
    body_text(doc,
        "Tu microservicio NUNCA debe esperar la respuesta de ms-auditoria antes de devolver "
        "su propia respuesta al cliente. El flujo correcto es:")
    numbered(doc, "Tu endpoint ejecuta su lógica de negocio.")
    numbered(doc, "Tu endpoint construye la respuesta para el cliente.")
    numbered(doc, "Tu endpoint devuelve la respuesta al cliente (HTTP 200/201/etc.).")
    numbered(doc, "En background (BackgroundTasks de FastAPI o asyncio.create_task), se envía el log a ms-auditoria.")
    numbered(doc, "ms-auditoria devuelve 202 Accepted. Esto no afecta al cliente de tu servicio.")

    callout(doc,
        "Jamás hagas await del envío a ms-auditoria dentro del handler principal de tu endpoint. "
        "Usa BackgroundTasks o asyncio.create_task para no bloquear.",
        kind='warning')

    heading2(doc, "Principio 2 — Si ms-auditoria falla, tu servicio continúa")
    body_text(doc,
        "El envío de logs debe estar rodeado de un bloque try/except. Si ms-auditoria devuelve "
        "un error, si está caído, o si hay un timeout de red, tu microservicio NO debe propagar "
        "ese error a su cliente. Solo debe registrar el error en sus propios logs de sistema.")
    code_block(doc, '''# Ejemplo correcto
async def send_audit_log(log_data: dict):
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            await client.post(
                f"{AUD_BASE_URL}/api/v1/logs",
                json=log_data,
                headers={"X-App-Token": AUD_APP_TOKEN}
            )
    except Exception as e:
        # Silenciar el error — ms-auditoria no puede tumbar tu servicio
        logger.warning(f"[AUDIT] No se pudo enviar log a ms-auditoria: {e}")''', lang_label='Python')

    heading2(doc, "Principio 3 — Registra TODAS las operaciones, no solo los errores")
    body_text(doc,
        "La regla transversal 6.6 establece que CADA operación genera un log. Esto incluye "
        "operaciones exitosas (200, 201), operaciones que devuelven 404, operaciones de consulta, "
        "operaciones de escritura, y errores internos (500). El objetivo es tener trazabilidad "
        "completa del sistema, no solo de sus fallos.")

    heading2(doc, "Principio 4 — Propaga el Request ID")
    body_text(doc,
        "Si tu microservicio recibe un header X-Request-ID de quien lo llamó, debes incluir ese "
        "mismo Request ID en el campo request_id del log que envías a ms-auditoria. Esto permite "
        "reconstruir toda la cadena de microservicios que participaron en una sola petición del "
        "usuario (traza distribuida). Si no recibes un X-Request-ID, genera uno propio con el "
        "formato: {CODIGO_MS}-{timestamp_unix_ms}-{6_chars_random}.")

    heading2(doc, "Principio 5 — No expongas datos sensibles en el campo detail")
    callout(doc,
        "El campo 'detail' del log es texto libre descriptivo. NUNCA debe contener contraseñas, "
        "tokens de sesión, datos de tarjetas de crédito, ni ningún dato personal sensible. "
        "Usa descripciones operacionales como: 'Matrícula MAT-2026-0042 actualizada a estado activo'.",
        kind='error')


def build_sec4_conexion(doc):
    add_page_break(doc)
    heading1(doc, "4. Información de Conexión y Endpoints")

    heading2(doc, "4.1 URL Base")
    info_table(doc,
        [
            ["Entorno", "URL Base"],
            ["Desarrollo local", "http://localhost:8019/api/v1"],
            ["Docker Compose (interno)", "http://ms-auditoria:8019/api/v1"],
            ["Producción / Staging", "https://api.erp-universitario.internal/aud/api/v1"],
        ],
        headers=["Entorno", "URL Base"],
        col_widths=[2.0, 4.5],
    )

    callout(doc,
        "En Docker Compose, usa el nombre del servicio como hostname: http://ms-auditoria:8019. "
        "No uses localhost entre contenedores.",
        kind='info')

    heading2(doc, "4.2 Endpoints disponibles para los microservicios emisores")
    body_text(doc,
        "Los microservicios del sistema (los 18 que integran con ms-auditoria) solo necesitan "
        "estos dos endpoints para enviar logs:")

    info_table(doc,
        [
            ["POST", "/api/v1/logs", "Enviar un log individual", "202 Accepted", "X-App-Token"],
            ["POST", "/api/v1/logs/batch", "Enviar un lote de logs (array)", "202 Accepted", "X-App-Token"],
            ["GET", "/api/v1/health", "Verificar estado de ms-auditoria", "200 healthy", "Ninguna"],
        ],
        headers=["Método", "Endpoint", "Descripción", "Respuesta", "Auth requerida"],
        col_widths=[0.8, 2.0, 2.5, 1.2, 1.2],
    )

    body_text(doc,
        "Los demás endpoints (consulta de logs, estadísticas, configuración de retención) son "
        "para uso exclusivo de administradores y analistas autenticados. Los microservicios del "
        "sistema no los necesitan para la integración básica.", italic=True)

    heading2(doc, "4.3 Timeouts recomendados")
    info_table(doc,
        [
            ["Timeout de conexión", "2 segundos", "Si ms-auditoria no responde en 2s, abortar y continuar"],
            ["Timeout de respuesta", "5 segundos", "Tiempo máximo total de la petición HTTP"],
            ["Reintentos", "0 (ninguno)", "No reintentar — el log puede perderse, el negocio no puede detenerse"],
        ],
        headers=["Parámetro", "Valor recomendado", "Razón"],
        col_widths=[1.8, 1.5, 3.0],
    )


def build_sec5_auth(doc):
    add_page_break(doc)
    heading1(doc, "5. Autenticación: El Token de Aplicación")

    body_text(doc,
        "Para enviar logs a ms-auditoria, cada microservicio se identifica con un Token de "
        "Aplicación. Este token es único por microservicio y se transmite "
        "en el header X-App-Token de cada petición.")

    callout(doc,
        "Comportamiento real del proyecto: en entorno development/testing, si no envías "
        "X-App-Token, el backend puede permitir la petición para facilitar pruebas locales. "
        "En producción el token sí es obligatorio.",
        kind='warning')

    heading2(doc, "5.1 ¿Cómo obtienes tu token de aplicación?")
    numbered(doc, "El administrador del sistema genera y registra el token de tu microservicio en ms-auditoria.")
    numbered(doc, "El token se almacena como variable de entorno en tu microservicio (en el gestor de secretos, Vault, o Kubernetes Secrets).")
    numbered(doc, "NUNCA incluyas el token en el código fuente ni en archivos de configuración versionados (.env no debe ir al repositorio).")
    numbered(doc, "La variable de entorno se llama convencionalmente: {CODIGO_MS}_APP_TOKEN (por ejemplo: RES_APP_TOKEN para ms-reservas).")

    heading2(doc, "5.2 Cómo incluir el token en la petición")
    body_text(doc, "El token se envía en el header HTTP X-App-Token en cada petición:")
    code_block(doc, '''POST /api/v1/logs HTTP/1.1
Host: ms-auditoria:8019
Content-Type: application/json
X-App-Token: {tu_token_de_aplicacion}
X-Request-ID: RES-1709302000-xyz789

{
  "timestamp": "2026-03-02T14:05:12Z",
  ...
}''', lang_label='HTTP Request')

    heading2(doc, "5.3 Configuración en Python")
    code_block(doc, '''# En tu archivo de configuración / settings.py
import os

# Token de aplicación de tu microservicio
# Cargado desde variable de entorno — NUNCA hardcodeado
APP_TOKEN = os.getenv("RES_APP_TOKEN")  # reemplaza RES por tu código

# URL base de ms-auditoria
AUD_BASE_URL = os.getenv("AUD_BASE_URL", "http://ms-auditoria:8019")

if not APP_TOKEN:
    raise RuntimeError("RES_APP_TOKEN no configurada. Revisar variables de entorno.")''', lang_label='Python — settings.py')

    callout(doc,
        "Si ms-auditoria responde 401 Unauthorized, significa que tu token de aplicación es "
        "inválido, está mal formateado, o no está registrado en ms-auditoria. Verifica con el "
        "administrador que tu token fue registrado correctamente.",
        kind='warning')

    heading2(doc, "5.4 Tabla de tokens por microservicio")
    body_text(doc,
        "Cada microservicio tiene su propio token. La siguiente tabla indica el nombre de "
        "variable de entorno que debe configurarse en cada uno:")

    info_table(doc,
        [
            ["ms-autenticacion", "AUTH_APP_TOKEN", "AUTH"],
            ["ms-usuarios",      "USR_APP_TOKEN",  "USR"],
            ["ms-roles",         "ROL_APP_TOKEN",  "ROL"],
            ["ms-inventario",    "INV_APP_TOKEN",  "INV"],
            ["ms-espacios",      "ESP_APP_TOKEN",  "ESP"],
            ["ms-reservas",      "RES_APP_TOKEN",  "RES"],
            ["ms-presupuesto",   "PRE_APP_TOKEN",  "PRE"],
            ["ms-gastos",        "GAS_APP_TOKEN",  "GAS"],
            ["ms-facturacion",   "FAC_APP_TOKEN",  "FAC"],
            ["ms-pedidos",       "PED_APP_TOKEN",  "PED"],
            ["ms-domicilios",    "DOM_APP_TOKEN",  "DOM"],
            ["ms-proveedores",   "PROV_APP_TOKEN", "PROV"],
            ["ms-programas",     "PROG_APP_TOKEN", "PROG"],
            ["ms-matriculas",    "MAT_APP_TOKEN",  "MAT"],
            ["ms-calificaciones","CAL_APP_TOKEN",  "CAL"],
            ["ms-horarios",      "HOR_APP_TOKEN",  "HOR"],
            ["ms-notificaciones","NOT_APP_TOKEN",  "NOT"],
            ["ms-reportes",      "REP_APP_TOKEN",  "REP"],
        ],
        headers=["Microservicio", "Variable de entorno", "Prefijo Request ID"],
        col_widths=[2.0, 2.0, 1.8],
    )


def build_sec6_estructura_log(doc):
    add_page_break(doc)
    heading1(doc, "6. Estructura del Registro de Log (Contrato de Datos)")

    body_text(doc,
        "Cada log que envías a ms-auditoria debe ser un objeto JSON con los siguientes campos. "
        "En el contrato actual, request_id y detail son opcionales (aunque se recomiendan para "
        "trazabilidad y diagnóstico).")

    heading2(doc, "6.1 Campos del log")
    info_table(doc,
        [
            ["timestamp",     "string (ISO 8601)", "Obligatorio", "Fecha y hora en que FINALIZÓ la operación. Formato: 2026-03-02T14:05:12Z. Siempre en UTC."],
            ["request_id",    "string (≤36 chars)", "Opcional", "Identificador de rastreo de la petición. Recomendado: propágalo desde el header X-Request-ID entrante, o genera uno propio con formato: {CODIGO}-{timestamp_ms}-{6chars}."],
            ["service_name",  "string", "Obligatorio", "Nombre de TU microservicio tal como está registrado en el sistema. Ejemplos: ms-reservas, ms-matriculas, ms-autenticacion."],
            ["functionality", "string", "Obligatorio", "Nombre de la funcionalidad/operación ejecutada. Usar snake_case. Ejemplos: crear_reserva, consultar_matriculas, login, actualizar_usuario."],
            ["method",        "string", "Obligatorio", "Método HTTP de la petición recibida por tu endpoint. En validación batch se aceptan: GET, POST, PUT, PATCH, DELETE, HEAD, OPTIONS."],
            ["response_code", "integer", "Obligatorio", "Código HTTP que tu endpoint devolvió al cliente. Rango válido: 100–599."],
            ["duration_ms",   "integer", "Obligatorio", "Duración de la operación en milisegundos (desde que llegó la petición hasta que enviaste la respuesta). Debe ser ≥ 0."],
            ["user_id",       "string | null", "Obligatorio*", "ID del usuario que realizó la operación. Si la operación no está asociada a un usuario (ej: health check, tarea interna), enviar null."],
            ["detail",        "string | null", "Opcional", "Descripción breve del resultado de la operación. Debe ser informativa pero NO contener datos sensibles. En el contrato actual soporta hasta 5000 caracteres."],
        ],
        headers=["Campo", "Tipo", "Requerido", "Descripción"],
        col_widths=[1.4, 1.4, 0.9, 3.3],
    )

    heading2(doc, "6.2 Ejemplo de log completo (ms-reservas)")
    code_block(doc, '''{
  "timestamp": "2026-03-02T14:05:12Z",
  "request_id": "RES-1709302000-xyz789",
  "service_name": "ms-reservas",
  "functionality": "crear_reserva",
  "method": "POST",
  "response_code": 201,
  "duration_ms": 312,
  "user_id": "usr-0002-uuid-docente",
  "detail": "Reserva creada exitosamente para espacio A-101, fecha 2026-03-10, usuario docente."
}''', lang_label='JSON — Log individual')

    heading2(doc, "6.3 Validaciones que aplica ms-auditoria")
    info_table(doc,
        [
            ["timestamp",     "Formato ISO 8601 válido. Si el formato es incorrecto → 422."],
            ["request_id",    "Opcional. Si se envía, longitud ≤ 36 caracteres."],
            ["service_name",  "No puede estar vacío. String no vacío."],
            ["functionality", "No puede estar vacío. String no vacío."],
            ["method",        "En batch se valida contra: GET, POST, PUT, PATCH, DELETE, HEAD, OPTIONS."],
            ["response_code", "Entero entre 100 y 599 (inclusive)."],
            ["duration_ms",   "Entero no negativo (≥ 0)."],
            ["user_id",       "Puede ser null. Si está presente, no puede ser string vacío."],
            ["detail",        "Opcional. Si se envía, máximo 5000 caracteres."],
        ],
        headers=["Campo", "Regla de validación"],
        col_widths=[1.5, 5.0],
    )

    callout(doc,
        "Si uno o más campos no pasan las validaciones, ms-auditoria responde 422 Unprocessable "
        "Entity con el detalle de qué campos fallaron. Revisa los logs de tu servicio para "
        "identificar el problema.",
        kind='warning')

    heading2(doc, "6.4 Errores frecuentes en el formato del log")
    info_table(doc,
        [
            ["timestamp con zona horaria incorrecta", '"2026-03-02 14:05:12"', '"2026-03-02T14:05:12Z"'],
            ["duration_ms negativo",                  '"-50"',                  '"312"'],
            ["response_code como string",             '"200"',                  '200 (entero)'],
            ["service_name vacío",                    '""',                     '"ms-reservas"'],
            ["detail con contraseña",                 '"Login con pass 1234"',  '"Login exitoso para usuario admin"'],
            ["user_id vacío en vez de null",          '""',                     'null'],
        ],
        headers=["Error", "Incorrecto ❌", "Correcto ✅"],
        col_widths=[2.4, 2.0, 2.0],
    )


def build_sec7_endpoint_individual(doc):
    add_page_break(doc)
    heading1(doc, "7. Endpoint Principal: POST /api/v1/logs")

    body_text(doc,
        "Este es el endpoint que usarás para la gran mayoría de tus operaciones. "
        "Envía un único registro de log por petición.")

    heading2(doc, "7.1 Especificación")
    info_table(doc,
        [
            ["Método HTTP",   "POST"],
            ["Endpoint",      "/api/v1/logs"],
            ["URL completa",  "http://ms-auditoria:8019/api/v1/logs"],
            ["Autenticación", "X-App-Token: {tu_token_de_aplicacion}"],
            ["Content-Type",  "application/json"],
            ["Respuesta éxito","202 Accepted — el log fue recibido y está en cola de almacenamiento"],
            ["Respuesta 401", "Token de aplicación inválido o ausente"],
            ["Respuesta 422", "Uno o más campos del log son inválidos"],
        ],
        headers=["Campo", "Valor"],
        col_widths=[2.0, 4.5],
    )

    heading2(doc, "7.2 Request completo")
    code_block(doc, '''POST /api/v1/logs HTTP/1.1
Host: ms-auditoria:8019
Content-Type: application/json
X-App-Token: token_de_aplicacion_registrado
X-Request-ID: RES-1709302000-xyz789

{
  "timestamp": "2026-03-02T14:05:12Z",
  "request_id": "RES-1709302000-xyz789",
  "service_name": "ms-reservas",
  "functionality": "crear_reserva",
  "method": "POST",
  "response_code": 201,
  "duration_ms": 312,
  "user_id": "usr-0002-uuid-docente",
  "detail": "Reserva creada para espacio A-101, fecha 2026-03-10."
}''', lang_label='HTTP Request')

    heading2(doc, "7.3 Respuesta exitosa (202)")
    code_block(doc, '''{
  "request_id": "AUD-1709302000-aud001",
  "success": true,
  "data": {
    "received": true,
    "log_request_id": "RES-1709302000-xyz789"
  },
  "message": "Registro de log recibido y en proceso de almacenamiento.",
  "timestamp": "2026-03-02T14:05:12Z"
}''', lang_label='JSON — Response 202')

    callout(doc,
        "El campo log_request_id en la respuesta es el request_id que enviaste en el body. "
        "Esto te permite confirmar que ms-auditoria recibió el log correcto.",
        kind='info')

    heading2(doc, "7.4 Respuesta de error 401 (token inválido)")
    code_block(doc, '''{
  "request_id": "AUD-1709302000-aud001",
  "success": false,
  "data": null,
  "message": "Token de aplicación inválido o no proporcionado.",
  "timestamp": "2026-03-02T14:05:12Z"
}''', lang_label='JSON — Response 401')

    heading2(doc, "7.5 Respuesta de error 422 (campos inválidos)")
    code_block(doc, '''{
  "request_id": "AUD-1709302000-aud001",
  "success": false,
  "data": {
    "invalid_fields": ["response_code", "duration_ms"],
    "detail": "response_code debe ser un entero entre 100 y 599; duration_ms debe ser un entero no negativo."
  },
  "message": "El cuerpo del log no cumple el formato requerido.",
  "timestamp": "2026-03-02T14:05:12Z"
}''', lang_label='JSON — Response 422')


def build_sec8_endpoint_batch(doc):
    add_page_break(doc)
    heading1(doc, "8. Endpoint para Lotes: POST /api/v1/logs/batch")

    body_text(doc,
        "Usa este endpoint cuando quieras enviar múltiples logs en una sola petición HTTP. "
        "Es más eficiente en términos de red cuando tienes operaciones que generan varios "
        "registros en un flujo corto (por ejemplo, un proceso batch que procesa 50 matrículas "
        "y quieres auditar cada una al final).")

    heading2(doc, "8.1 Cuándo usar batch vs. individual")
    info_table(doc,
        [
            ["Operación normal de endpoint (1 petición → 1 log)", "✅ Usa /logs (individual)", "Más simple, menos overhead"],
            ["Proceso batch / importación masiva (N items → N logs)", "✅ Usa /logs/batch", "Reduce N peticiones a 1"],
            ["Flush periódico de un buffer de logs acumulados", "✅ Usa /logs/batch", "Agrupa en un solo request"],
            ["Solo tienes 1-2 logs juntos", "✅ Usa /logs (individual) x2", "El overhead de batch no vale la pena"],
        ],
        headers=["Escenario", "Recomendación", "Razón"],
        col_widths=[2.5, 2.0, 2.0],
    )

    heading2(doc, "8.2 Especificación")
    info_table(doc,
        [
            ["Método HTTP",    "POST"],
            ["Endpoint",       "/api/v1/logs/batch"],
            ["Autenticación",  "X-App-Token: {tu_token_de_aplicacion}"],
            ["Content-Type",   "application/json"],
            ["Body requerido", "Objeto JSON con campo 'logs': arreglo de objetos log (mínimo 1 elemento)"],
            ["Respuesta éxito", "202 Accepted — incluye conteo de aceptados y rechazados por separado"],
        ],
        headers=["Campo", "Valor"],
        col_widths=[2.0, 4.5],
    )

    heading2(doc, "8.3 Request completo")
    code_block(doc, '''POST /api/v1/logs/batch HTTP/1.1
Host: ms-auditoria:8019
Content-Type: application/json
X-App-Token: token_de_aplicacion_registrado

{
  "logs": [
    {
      "timestamp": "2026-03-02T14:05:10Z",
      "request_id": "MAT-1709302000-001",
      "service_name": "ms-matriculas",
      "functionality": "procesar_lote_matriculas",
      "method": "POST",
      "response_code": 201,
      "duration_ms": 95,
      "user_id": "usr-0001-uuid-admin",
      "detail": "Matrícula MAT-2026-0041 creada para estudiante USR-0042."
    },
    {
      "timestamp": "2026-03-02T14:05:11Z",
      "request_id": "MAT-1709302000-002",
      "service_name": "ms-matriculas",
      "functionality": "procesar_lote_matriculas",
      "method": "POST",
      "response_code": 201,
      "duration_ms": 88,
      "user_id": "usr-0001-uuid-admin",
      "detail": "Matrícula MAT-2026-0042 creada para estudiante USR-0043."
    }
  ]
}''', lang_label='HTTP Request — Batch')

    heading2(doc, "8.4 Respuesta exitosa (202)")
    code_block(doc, '''{
  "request_id": "AUD-1709302000-aud002",
  "success": true,
  "data": {
    "received_count": 2,
    "accepted_count": 2,
    "rejected_count": 0,
    "rejected_indices": []
  },
  "message": "Lote de logs recibido. 2 registros aceptados para almacenamiento.",
  "timestamp": "2026-03-02T14:05:12Z"
}''', lang_label='JSON — Response 202 OK')

    heading2(doc, "8.5 Respuesta con registros rechazados (202 parcial)")
    body_text(doc,
        "Si algunos registros del lote tienen campos inválidos, ms-auditoria igualmente "
        "responde 202 pero informa cuáles fueron rechazados:")
    code_block(doc, '''{
  "request_id": "AUD-1709302000-aud003",
  "success": true,
  "data": {
    "received_count": 3,
    "accepted_count": 2,
    "rejected_count": 1,
    "rejected_indices": [
      {
        "index": 2,
        "reason": "duration_ms debe ser un entero no negativo."
      }
    ]
  },
  "message": "Lote de logs recibido. 2 de 3 registros aceptados para almacenamiento.",
  "timestamp": "2026-03-02T14:05:12Z"
}''', lang_label='JSON — Response 202 con rechazo parcial')

    callout(doc,
        "Revisa rejected_indices en tu log de sistema cuando haya rechazos. El índice 2 significa "
        "el tercer elemento del arreglo (índice 0-based). Corrige el log en tu servicio y reenvíalo "
        "si es crítico para tu auditoría.",
        kind='warning')


def build_sec9_errores(doc):
    add_page_break(doc)
    heading1(doc, "9. Manejo de Errores y Respuestas")

    heading2(doc, "9.1 Tabla de códigos de respuesta")
    info_table(doc,
        [
            ["202 Accepted",               "✅ Éxito",   "Log recibido correctamente. El almacenamiento ocurre en segundo plano."],
            ["401 Unauthorized",           "❌ Error",   "X-App-Token ausente, malformado o no registrado. Verifica tu token de aplicación."],
            ["422 Unprocessable Entity",   "❌ Error",   "Uno o más campos del log tienen formato inválido. Ver campo data.invalid_fields en la respuesta."],
            ["500 Internal Server Error",  "❌ Error",   "Error interno de ms-auditoria. Registra el error en tu sistema y continúa operando."],
            ["503 Service Unavailable",    "❌ Error",   "ms-auditoria no está disponible. Tu servicio debe continuar sin interrupciones."],
        ],
        headers=["Código", "Tipo", "Qué hacer"],
        col_widths=[1.5, 0.8, 4.2],
    )

    heading2(doc, "9.2 Estructura estándar de respuesta de ms-auditoria")
    body_text(doc,
        "Todos los endpoints de ms-auditoria usan la misma estructura de respuesta JSON:")
    code_block(doc, '''{
  "request_id": "AUD-{timestamp}-{id}",   // Request ID de la operación en ms-auditoria
  "success": true | false,                 // Indicador de éxito
  "data": { ... } | null,                 // Datos de la respuesta (null si hay error)
  "message": "Descripción del resultado", // Mensaje descriptivo
  "timestamp": "2026-03-02T14:05:12Z"    // Fecha/hora de la respuesta en ms-auditoria
}''', lang_label='JSON — Estructura de respuesta estándar')

    heading2(doc, "9.3 Comportamiento recomendado por código de respuesta")
    info_table(doc,
        [
            ["202", "No hacer nada especial. Opción: loguear en DEBUG que el audit fue enviado."],
            ["401", "Loguear ERROR en tu sistema con detalle. Revisar variable de entorno del token. NO reintentar automáticamente."],
            ["422", "Loguear ERROR en tu sistema con el detalle de campos inválidos. Corregir el log en el código fuente. NO reintentar con el mismo payload."],
            ["500 / 503", "Loguear WARNING en tu sistema. Tu servicio continúa operando normalmente. El log se pierde (aceptable por diseño)."],
            ["Timeout / red", "Loguear WARNING en tu sistema. Tu servicio continúa operando normalmente. El log se pierde (aceptable por diseño)."],
        ],
        headers=["Código", "Qué hacer en tu servicio"],
        col_widths=[0.8, 5.7],
    )


def build_sec10_implementacion(doc):
    add_page_break(doc)
    heading1(doc, "10. Patrón de Implementación Recomendado (Python / FastAPI)")

    body_text(doc,
        "Esta sección contiene el código completo y listo para usar que debes agregar a tu "
        "microservicio para integrar con ms-auditoria. El patrón está diseñado para Python "
        "con FastAPI y httpx (cliente HTTP asíncrono).")

    heading2(doc, "10.1 Instalación de dependencias")
    code_block(doc, '''# En requirements.txt de tu microservicio
httpx>=0.27.0        # Cliente HTTP asíncrono
python-dotenv>=1.0.0 # Variables de entorno''', lang_label='requirements.txt')

    heading2(doc, "10.2 Módulo de cliente de auditoría (audit_client.py)")
    body_text(doc,
        "Crea este archivo en tu proyecto. Es el único punto de contacto con ms-auditoria:")
    code_block(doc, '''# audit_client.py
"""
Cliente de integración con ms-auditoria.
Uso:
    from audit_client import send_audit_log
    await send_audit_log(request, "crear_reserva", "POST", 201, 312,
                         "usr-0002-uuid-docente",
                         "Reserva creada para espacio A-101.")
"""
import os
import asyncio
import logging
import time
from datetime import datetime, timezone
from typing import Optional
import httpx

logger = logging.getLogger(__name__)

# ── Configuración ──────────────────────────────────────────────────────────
AUD_BASE_URL  = os.getenv("AUD_BASE_URL", "http://ms-auditoria:8019")
AUD_ENDPOINT  = f"{AUD_BASE_URL}/api/v1/logs"
APP_TOKEN     = os.getenv("APP_TOKEN")       # Token de aplicación de tu MS
SERVICE_NAME  = os.getenv("SERVICE_NAME")    # Ej: "ms-reservas"
TIMEOUT_SECS  = float(os.getenv("AUD_TIMEOUT", "5.0"))


def _get_request_id(request) -> str:
    """Extrae X-Request-ID del header entrante o genera uno nuevo."""
    rid = None
    if request and hasattr(request, 'headers'):
        rid = request.headers.get("x-request-id")
    if not rid:
        import random, string
        ts  = int(time.time() * 1000)
        sfx = ''.join(random.choices(string.ascii_lowercase + string.digits, k=6))
        prefix = (SERVICE_NAME or "MS").upper().replace("MS-", "")
        rid = f"{prefix}-{ts}-{sfx}"
    return rid


async def _post_log(payload: dict) -> None:
    """Envía el log a ms-auditoria. Silencia cualquier error."""
    if not APP_TOKEN:
        logger.warning("[AUDIT] APP_TOKEN no configurado. Log no enviado.")
        return
    headers = {
        "Content-Type": "application/json",
        "X-App-Token": APP_TOKEN,
        "X-Request-ID": payload.get("request_id", ""),
    }
    try:
        async with httpx.AsyncClient(timeout=TIMEOUT_SECS) as client:
            response = await client.post(AUD_ENDPOINT, json=payload, headers=headers)
            if response.status_code == 401:
                logger.error("[AUDIT] Token inválido (401). Revisar APP_TOKEN.")
            elif response.status_code == 422:
                logger.error(f"[AUDIT] Payload inválido (422): {response.text}")
            elif response.status_code not in (200, 201, 202):
                logger.warning(f"[AUDIT] Respuesta inesperada: {response.status_code}")
    except httpx.TimeoutException:
        logger.warning("[AUDIT] Timeout al conectar con ms-auditoria.")
    except Exception as exc:
        logger.warning(f"[AUDIT] Error al enviar log: {exc}")


async def send_audit_log(
    request,                    # FastAPI Request object (o None)
    functionality: str,
    method: str,
    response_code: int,
    duration_ms: int,
    user_id: Optional[str],
    detail: str,
) -> None:
    """
    Construye y envía un log a ms-auditoria de forma asíncrona.
    Esta función NO bloquea — usa asyncio.create_task internamente.
    """
    payload = {
        "timestamp":     datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        "request_id":    _get_request_id(request),
        "service_name":  SERVICE_NAME or "ms-desconocido",
        "functionality": functionality,
        "method":        method.upper(),
        "response_code": response_code,
        "duration_ms":   max(0, duration_ms),
        "user_id":       user_id,
        "detail":        detail,
    }
    # Fire-and-forget: no espera la respuesta
    asyncio.create_task(_post_log(payload))''', lang_label='Python — audit_client.py')

    heading2(doc, "10.3 Uso en tus endpoints FastAPI")
    code_block(doc, '''# En tu archivo de router / endpoint
import time
from fastapi import APIRouter, Request, BackgroundTasks
from audit_client import send_audit_log

router = APIRouter()

@router.post("/reservas", status_code=201)
async def crear_reserva(
    request: Request,
    payload: ReservaCreate,
    background_tasks: BackgroundTasks
):
    start_time = time.time()

    # ── Tu lógica de negocio aquí ─────────────────────────────────────
    nueva_reserva = await reserva_service.crear(payload)
    # ─────────────────────────────────────────────────────────────────

    duration_ms = int((time.time() - start_time) * 1000)

    # ── Envío asíncrono del log (no bloquea la respuesta) ─────────────
    background_tasks.add_task(
        send_audit_log,
        request     = request,
        functionality = "crear_reserva",
        method      = "POST",
        response_code = 201,
        duration_ms = duration_ms,
        user_id     = request.state.user_id,  # del middleware de auth
        detail      = f"Reserva creada para espacio {nueva_reserva.espacio_id}, "
                      f"fecha {nueva_reserva.fecha}."
    )
    # ─────────────────────────────────────────────────────────────────

    return nueva_reserva''', lang_label='Python — Endpoint de ejemplo (ms-reservas)')

    heading2(doc, "10.4 Variables de entorno necesarias en tu microservicio")
    code_block(doc, '''# .env de tu microservicio (NO subir al repositorio)
APP_TOKEN=tu_token_de_aplicacion_cifrado
SERVICE_NAME=ms-reservas           # Nombre exacto de tu microservicio
AUD_BASE_URL=http://ms-auditoria:8019
AUD_TIMEOUT=5.0                    # Segundos de timeout para ms-auditoria''', lang_label='.env — Variables de entorno')

    callout(doc,
        "Usa BackgroundTasks de FastAPI para el envío del log en lugar de asyncio.create_task "
        "cuando estés dentro de un endpoint FastAPI. BackgroundTasks garantiza que la tarea se "
        "ejecuta después de enviar la respuesta y está manejada por el ciclo de vida de FastAPI.",
        kind='info')

    heading2(doc, "10.5 Middleware para medir duración automáticamente")
    body_text(doc,
        "Para no tener que calcular la duración manualmente en cada endpoint, puedes agregar "
        "un middleware que mida el tiempo de cada petición y lo almacene en request.state:")
    code_block(doc, '''# middleware.py
import time
from fastapi import Request

async def timing_middleware(request: Request, call_next):
    """Middleware que mide la duración de cada petición."""
    request.state.start_time = time.time()
    response = await call_next(request)
    duration_ms = int((time.time() - request.state.start_time) * 1000)
    response.headers["X-Duration-Ms"] = str(duration_ms)
    request.state.duration_ms = duration_ms
    return response

# En main.py:
# app.middleware("http")(timing_middleware)

# En tu endpoint, accede a request.state.duration_ms:
background_tasks.add_task(
    send_audit_log,
    request=request,
    functionality="crear_reserva",
    method="POST",
    response_code=201,
    duration_ms=request.state.duration_ms,  # Automático del middleware
    user_id=request.state.user_id,
    detail="Reserva creada."
)''', lang_label='Python — Middleware de timing')


def build_sec11_request_id(doc):
    add_page_break(doc)
    heading1(doc, "11. Request ID: Trazabilidad Distribuida")

    body_text(doc,
        "El Request ID es el mecanismo que permite reconstruir toda la cadena de llamadas "
        "entre microservicios para una misma petición del usuario. Si el usuario hace click "
        "en 'Crear Reserva' y eso desencadena llamadas a ms-autenticacion, ms-roles y "
        "ms-reservas, todos deben usar el mismo Request ID para que ms-auditoria pueda "
        "reconstruir la traza completa.")

    heading2(doc, "11.1 Formato del Request ID")
    code_block(doc, '''{CODIGO_MS}-{timestamp_unix_ms}-{6_chars_random}

Ejemplos:
  RES-1709302000000-a3f7c2    (ms-reservas)
  MAT-1709302000000-b8d4e1    (ms-matriculas)
  AUTH-1709302000000-c5f2a3   (ms-autenticacion)
  AUD-1709302000000-d9e6b7    (ms-auditoria, generado internamente)''', lang_label='Formato Request ID')

    heading2(doc, "11.2 Reglas de propagación")
    numbered(doc, "Cuando recibes una petición con header X-Request-ID → úsalo como request_id en tu log.")
    numbered(doc, "Cuando recibes una petición SIN header X-Request-ID → genera uno nuevo con tu prefijo.")
    numbered(doc, "Cuando llamas a OTRO microservicio → pasa el X-Request-ID en el header de tu petición saliente.")
    numbered(doc, "El campo request_id del log es siempre el X-Request-ID de la petición que llegó a TU endpoint.")

    heading2(doc, "11.3 Ejemplo de propagación en un flujo de 3 servicios")
    code_block(doc, '''# Escenario: Usuario crea una reserva

# 1. Frontend → ms-reservas
# (el frontend genera el Request ID inicial)
POST /reservas
X-Request-ID: FRONT-1709302000000-aa1234

# 2. ms-reservas verifica espacio en ms-espacios
GET /espacios/A-101/disponibilidad
X-Request-ID: FRONT-1709302000000-aa1234  ← MISMO ID propagado

# 3. ms-reservas registra log en ms-auditoria
POST /api/v1/logs
{
  "request_id": "FRONT-1709302000000-aa1234",  ← MISMO ID en el log
  "service_name": "ms-reservas",
  "functionality": "crear_reserva",
  ...
}

# Resultado: la traza completa visible en ms-auditoria filtrando por
# request_id = "FRONT-1709302000000-aa1234"''', lang_label='Ejemplo de propagación de Request ID')

    heading2(doc, "11.4 Código para propagación del Request ID")
    code_block(doc, '''# En tu audit_client.py ya tienes _get_request_id().
# Para propagar en llamadas salientes a otros microservicios:

async def call_otro_microservicio(request: Request, url: str, payload: dict):
    """Ejemplo de llamada saliente con propagación de Request ID."""
    rid = request.headers.get("x-request-id", "")
    headers = {
        "X-Request-ID": rid,
        "Authorization": f"Bearer {request.headers.get("authorization", "")}",
        "Content-Type": "application/json",
    }
    async with httpx.AsyncClient(timeout=10.0) as client:
        response = await client.get(url, headers=headers)
    return response''', lang_label='Python — Propagación en llamadas salientes')


def build_sec12_health(doc):
    add_page_break(doc)
    heading1(doc, "12. Health Check")

    body_text(doc,
        "Puedes verificar si ms-auditoria está disponible haciendo un GET al endpoint "
        "de health check. No requiere ninguna autenticación.")

    code_block(doc, '''GET /api/v1/health HTTP/1.1
Host: ms-auditoria:8019''', lang_label='HTTP Request')

    body_text(doc, "Respuesta cuando el servicio está saludable (200):")
    code_block(doc, '''{
  "status": "healthy",
    "service": "ms-auditoria",
  "version": "1.0.0",
  "timestamp": "2026-03-02T16:10:00Z",
  "components": {
    "database": {
      "status": "healthy",
      "latency_ms": 3
    }
  }
}''', lang_label='JSON — 200 Healthy')

    body_text(doc, "Respuesta cuando hay problemas (503):")
    code_block(doc, '''{
  "status": "unhealthy",
    "service": "ms-auditoria",
  "version": "1.0.0",
  "timestamp": "2026-03-02T16:10:00Z",
  "components": {
    "database": {
      "status": "unhealthy",
            "error": "Connection timeout after 5000ms"
    }
  }
}''', lang_label='JSON — 503 Unhealthy')

    callout(doc,
        "Puedes usar el health check en el startup de tu microservicio para verificar que "
        "ms-auditoria está disponible, pero recuerda: aunque ms-auditoria esté unhealthy, "
        "tu servicio debe arrancar y operar con normalidad (el envío de logs fallará "
        "silenciosamente).",
        kind='info')


def build_sec13_ejemplos_por_ms(doc):
    add_page_break(doc)
    heading1(doc, "13. Ejemplos de Integración por Microservicio (18 servicios)")

    body_text(doc,
        "Esta sección presenta ejemplos concretos de los logs que cada microservicio del "
        "sistema debe enviar a ms-auditoria. Para cada servicio se muestran 2-3 operaciones "
        "representativas con su log completo.")

    callout(doc,
        "Estos ejemplos muestran el payload exacto del body JSON. El header X-App-Token "
        "es obligatorio en producción y varía por microservicio (ver Sección 5).",
        kind='info')

    # ── Datos de los 18 microservicios ───────────────────────────────────────
    microservicios = [
        {
            "nombre": "ms-autenticacion [AUTH]",
            "desc": "Gestiona la autenticación de usuarios, sesiones y tokens.",
            "prefix": "AUTH",
            "ops": [
                {
                    "func": "login",
                    "method": "POST",
                    "code": 200,
                    "ms": 145,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Inicio de sesión exitoso para usuario administrador. Sesión creada."
                },
                {
                    "func": "login",
                    "method": "POST",
                    "code": 401,
                    "ms": 55,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Intento de login fallido para usuario admin. Contraseña incorrecta."
                },
                {
                    "func": "logout",
                    "method": "POST",
                    "code": 200,
                    "ms": 32,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Sesión cerrada correctamente para usuario administrador."
                },
            ]
        },
        {
            "nombre": "ms-usuarios [USR]",
            "desc": "Gestiona el registro y perfil de usuarios del sistema.",
            "prefix": "USR",
            "ops": [
                {
                    "func": "crear_usuario",
                    "method": "POST",
                    "code": 201,
                    "ms": 280,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Usuario USR-0055 creado con rol estudiante. Email: estudiante@univalle.edu.co."
                },
                {
                    "func": "actualizar_usuario",
                    "method": "PATCH",
                    "code": 200,
                    "ms": 190,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Datos de usuario USR-0055 actualizados: teléfono y dirección."
                },
                {
                    "func": "consultar_usuario",
                    "method": "GET",
                    "code": 404,
                    "ms": 45,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Usuario USR-9999 no encontrado en el sistema."
                },
            ]
        },
        {
            "nombre": "ms-roles [ROL]",
            "desc": "Gestiona los roles y permisos por funcionalidad.",
            "prefix": "ROL",
            "ops": [
                {
                    "func": "asignar_rol",
                    "method": "POST",
                    "code": 201,
                    "ms": 220,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Rol DOCENTE asignado a usuario USR-0042. Asignado por administrador."
                },
                {
                    "func": "verificar_permisos",
                    "method": "GET",
                    "code": 200,
                    "ms": 35,
                    "uid": "usr-0042-uuid-docente",
                    "detail": "Permiso RES_CREAR_RESERVA verificado para usuario USR-0042: autorizado."
                },
                {
                    "func": "revocar_rol",
                    "method": "DELETE",
                    "code": 200,
                    "ms": 175,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Rol ESTUDIANTE revocado al usuario USR-0030 por vencimiento de matrícula."
                },
            ]
        },
        {
            "nombre": "ms-inventario [INV]",
            "desc": "Gestiona activos físicos e inventario de la institución.",
            "prefix": "INV",
            "ops": [
                {
                    "func": "registrar_activo",
                    "method": "POST",
                    "code": 201,
                    "ms": 310,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Activo INV-2026-0101 registrado: Computador portátil Dell, Bodega A."
                },
                {
                    "func": "consultar_activos",
                    "method": "GET",
                    "code": 200,
                    "ms": 88,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Consulta de activos filtrada por categoría=equipos. Total 47 registros."
                },
                {
                    "func": "actualizar_estado_activo",
                    "method": "PATCH",
                    "code": 200,
                    "ms": 155,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Activo INV-2026-0101 cambiado a estado en_mantenimiento."
                },
            ]
        },
        {
            "nombre": "ms-espacios [ESP]",
            "desc": "Gestiona los espacios físicos (aulas, laboratorios, salas).",
            "prefix": "ESP",
            "ops": [
                {
                    "func": "crear_espacio",
                    "method": "POST",
                    "code": 201,
                    "ms": 265,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Espacio A-101 creado: Aula de clase, capacidad 35 personas, Edificio A."
                },
                {
                    "func": "consultar_disponibilidad",
                    "method": "GET",
                    "code": 200,
                    "ms": 95,
                    "uid": "usr-0042-uuid-docente",
                    "detail": "Disponibilidad consultada para espacio A-101 en fecha 2026-03-10. Disponible."
                },
                {
                    "func": "actualizar_espacio",
                    "method": "PATCH",
                    "code": 200,
                    "ms": 180,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Espacio A-101 actualizado: capacidad cambiada de 35 a 30 personas."
                },
            ]
        },
        {
            "nombre": "ms-reservas [RES]",
            "desc": "Gestiona las reservas de espacios físicos.",
            "prefix": "RES",
            "ops": [
                {
                    "func": "crear_reserva",
                    "method": "POST",
                    "code": 201,
                    "ms": 312,
                    "uid": "usr-0042-uuid-docente",
                    "detail": "Reserva RES-2026-0089 creada para espacio A-101, fecha 2026-03-10, 08:00-10:00."
                },
                {
                    "func": "cancelar_reserva",
                    "method": "DELETE",
                    "code": 200,
                    "ms": 145,
                    "uid": "usr-0042-uuid-docente",
                    "detail": "Reserva RES-2026-0089 cancelada por el docente con 48h de anticipación."
                },
                {
                    "func": "consultar_reservas",
                    "method": "GET",
                    "code": 200,
                    "ms": 78,
                    "uid": "usr-0042-uuid-docente",
                    "detail": "Consulta de reservas para espacio A-101, semana del 2026-03-10. Total 5 registros."
                },
            ]
        },
        {
            "nombre": "ms-presupuesto [PRE]",
            "desc": "Gestiona los presupuestos por área y periodo.",
            "prefix": "PRE",
            "ops": [
                {
                    "func": "crear_presupuesto",
                    "method": "POST",
                    "code": 201,
                    "ms": 290,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Presupuesto PRE-2026-001 creado para área Académica, periodo 2026-1. Monto: $50.000.000."
                },
                {
                    "func": "consultar_ejecucion",
                    "method": "GET",
                    "code": 200,
                    "ms": 112,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Ejecución presupuestal consultada para área Académica, periodo 2026-1. Ejecutado: 35%."
                },
            ]
        },
        {
            "nombre": "ms-gastos [GAS]",
            "desc": "Gestiona los gastos y novedades del área financiera.",
            "prefix": "GAS",
            "ops": [
                {
                    "func": "registrar_gasto",
                    "method": "POST",
                    "code": 201,
                    "ms": 335,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Gasto GAS-2026-0412 registrado: Compra de resmas de papel, $340.000. Área: Académica."
                },
                {
                    "func": "aprobar_gasto",
                    "method": "PATCH",
                    "code": 200,
                    "ms": 178,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Gasto GAS-2026-0412 aprobado por administrador. Presupuesto descontado."
                },
            ]
        },
        {
            "nombre": "ms-facturacion [FAC]",
            "desc": "Gestiona la facturación y cobros de servicios académicos.",
            "prefix": "FAC",
            "ops": [
                {
                    "func": "generar_factura",
                    "method": "POST",
                    "code": 201,
                    "ms": 445,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Factura FAC-2026-0211 generada para estudiante USR-0030. Concepto: matrícula 2026-1."
                },
                {
                    "func": "registrar_pago",
                    "method": "POST",
                    "code": 200,
                    "ms": 289,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Pago registrado para factura FAC-2026-0211. Monto: $2.500.000. Estado: pagada."
                },
            ]
        },
        {
            "nombre": "ms-pedidos [PED]",
            "desc": "Gestiona los pedidos de materiales y suministros.",
            "prefix": "PED",
            "ops": [
                {
                    "func": "crear_pedido",
                    "method": "POST",
                    "code": 201,
                    "ms": 315,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Pedido PED-2026-0078 creado a proveedor PROV-005. 200 resmas papel. Urgente."
                },
                {
                    "func": "actualizar_estado_pedido",
                    "method": "PATCH",
                    "code": 200,
                    "ms": 165,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Pedido PED-2026-0078 actualizado a estado en_tránsito. Fecha estimada: 2026-03-15."
                },
            ]
        },
        {
            "nombre": "ms-domicilios [DOM]",
            "desc": "Gestiona las entregas y domicilios de pedidos.",
            "prefix": "DOM",
            "ops": [
                {
                    "func": "registrar_entrega",
                    "method": "POST",
                    "code": 201,
                    "ms": 275,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Entrega DOM-2026-0031 registrada para pedido PED-2026-0078. Recibido en bodega A."
                },
                {
                    "func": "consultar_entregas_pendientes",
                    "method": "GET",
                    "code": 200,
                    "ms": 90,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Consulta de entregas pendientes. Total 3 entregas por confirmar."
                },
            ]
        },
        {
            "nombre": "ms-proveedores [PROV]",
            "desc": "Gestiona el registro y catálogo de proveedores.",
            "prefix": "PROV",
            "ops": [
                {
                    "func": "registrar_proveedor",
                    "method": "POST",
                    "code": 201,
                    "ms": 300,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Proveedor PROV-010 registrado: Papelería El Estudiante, NIT 900.123.456-1."
                },
                {
                    "func": "actualizar_proveedor",
                    "method": "PATCH",
                    "code": 200,
                    "ms": 170,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Proveedor PROV-010 actualizado: nuevo contacto y teléfono registrados."
                },
            ]
        },
        {
            "nombre": "ms-programas [PROG]",
            "desc": "Gestiona los programas académicos y su estructura curricular.",
            "prefix": "PROG",
            "ops": [
                {
                    "func": "crear_programa",
                    "method": "POST",
                    "code": 201,
                    "ms": 420,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Programa PROG-001 creado: Ingeniería de Sistemas, duración 10 semestres."
                },
                {
                    "func": "consultar_programas",
                    "method": "GET",
                    "code": 200,
                    "ms": 65,
                    "uid": "usr-0042-uuid-docente",
                    "detail": "Consulta de programas activos. Total 5 programas disponibles."
                },
                {
                    "func": "actualizar_pensum",
                    "method": "PATCH",
                    "code": 200,
                    "ms": 385,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Pensum de programa PROG-001 actualizado: materia Arquitectura de Software agregada."
                },
            ]
        },
        {
            "nombre": "ms-matriculas [MAT]",
            "desc": "Gestiona el proceso de matrículas de estudiantes.",
            "prefix": "MAT",
            "ops": [
                {
                    "func": "crear_matricula",
                    "method": "POST",
                    "code": 201,
                    "ms": 480,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Matrícula MAT-2026-0042 creada para estudiante USR-0030, programa PROG-001, periodo 2026-1."
                },
                {
                    "func": "actualizar_estado_matricula",
                    "method": "PATCH",
                    "code": 200,
                    "ms": 210,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Matrícula MAT-2026-0042 cambiada a estado activa. Pago confirmado."
                },
                {
                    "func": "consultar_matriculas",
                    "method": "GET",
                    "code": 200,
                    "ms": 95,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Consulta de matrículas periodo 2026-1. Total 128 matrículas activas."
                },
            ]
        },
        {
            "nombre": "ms-calificaciones [CAL]",
            "desc": "Gestiona las calificaciones y notas de los estudiantes.",
            "prefix": "CAL",
            "ops": [
                {
                    "func": "registrar_calificacion",
                    "method": "POST",
                    "code": 201,
                    "ms": 295,
                    "uid": "usr-0042-uuid-docente",
                    "detail": "Calificación registrada para estudiante USR-0030, materia Matemáticas, nota: 4.2."
                },
                {
                    "func": "consultar_calificaciones_estudiante",
                    "method": "GET",
                    "code": 200,
                    "ms": 110,
                    "uid": "usr-0030-uuid-estudiante",
                    "detail": "Calificaciones consultadas para periodo 2026-1. Total 6 materias, promedio 3.9."
                },
                {
                    "func": "cerrar_periodo_calificaciones",
                    "method": "POST",
                    "code": 200,
                    "ms": 1250,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Periodo 2026-1 cerrado. 128 estudiantes procesados. Promedios calculados."
                },
            ]
        },
        {
            "nombre": "ms-horarios [HOR]",
            "desc": "Gestiona los horarios de clases y asignación de docentes.",
            "prefix": "HOR",
            "ops": [
                {
                    "func": "crear_horario",
                    "method": "POST",
                    "code": 201,
                    "ms": 365,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Horario creado para materia Arquitectura de Software, docente USR-0042, Lun-Mié 08:00-10:00, A-101."
                },
                {
                    "func": "consultar_horario_docente",
                    "method": "GET",
                    "code": 200,
                    "ms": 85,
                    "uid": "usr-0042-uuid-docente",
                    "detail": "Horario del docente USR-0042 consultado. 4 materias asignadas para 2026-1."
                },
            ]
        },
        {
            "nombre": "ms-notificaciones [NOT]",
            "desc": "Gestiona el envío de notificaciones y mensajes del sistema.",
            "prefix": "NOT",
            "ops": [
                {
                    "func": "enviar_notificacion",
                    "method": "POST",
                    "code": 202,
                    "ms": 125,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Notificación enviada a 128 estudiantes: recordatorio de pago de matrícula 2026-1."
                },
                {
                    "func": "enviar_notificacion",
                    "method": "POST",
                    "code": 500,
                    "ms": 5001,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Error al enviar notificación a estudiante USR-0030. Servicio SMTP no disponible."
                },
                {
                    "func": "consultar_notificaciones",
                    "method": "GET",
                    "code": 200,
                    "ms": 72,
                    "uid": "usr-0030-uuid-estudiante",
                    "detail": "Notificaciones del usuario USR-0030 consultadas. Total 3 notificaciones no leídas."
                },
            ]
        },
        {
            "nombre": "ms-reportes [REP]",
            "desc": "Genera reportes y estadísticas consolidadas del sistema.",
            "prefix": "REP",
            "ops": [
                {
                    "func": "generar_reporte_academico",
                    "method": "POST",
                    "code": 202,
                    "ms": 2150,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Reporte académico periodo 2026-1 generado. 128 estudiantes incluidos. Archivo: reporte_2026-1.pdf."
                },
                {
                    "func": "generar_reporte_financiero",
                    "method": "POST",
                    "code": 202,
                    "ms": 1880,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Reporte financiero mensual Febrero 2026 generado. Total facturado: $320.000.000."
                },
                {
                    "func": "consultar_reportes",
                    "method": "GET",
                    "code": 200,
                    "ms": 95,
                    "uid": "usr-0001-uuid-admin",
                    "detail": "Reportes disponibles consultados. Total 12 reportes generados en Febrero 2026."
                },
            ]
        },
    ]

    for ms_data in microservicios:
        heading2(doc, f"13.x {ms_data['nombre']}")
        body_text(doc, ms_data['desc'], italic=True)

        for i, op in enumerate(ms_data['ops'], 1):
            heading3(doc, f"Operación: {op['func']} — {op['method']} {op['code']}")

            ts = f"2026-03-02T14:0{i}:{10+i*5}Z"
            rid = f"{ms_data['prefix']}-17093020{i:02d}000-{op['func'][:3]}{i:03d}"
            code_block(doc, f'''{{\n  "timestamp": "{ts}",\n  "request_id": "{rid}",\n  "service_name": "{ms_data['nombre'].split(' ')[0]}",\n  "functionality": "{op['func']}",\n  "method": "{op['method']}",\n  "response_code": {op['code']},\n  "duration_ms": {op['ms']},\n  "user_id": "{op['uid']}",\n  "detail": "{op['detail']}"\n}}''',
                       lang_label=f'JSON — Log de {ms_data["nombre"].split(" ")[0]}')

        add_horizontal_rule(doc)


def build_sec14_troubleshooting(doc):
    add_page_break(doc)
    heading1(doc, "14. Guía de Resolución de Problemas")

    heading2(doc, "Problema 1 — ms-auditoria devuelve 401 Unauthorized")
    callout(doc, "Tu token de aplicación no está siendo reconocido.", kind='error')
    numbered(doc, "Verifica que la variable de entorno APP_TOKEN está configurada en tu microservicio.")
    numbered(doc, "Verifica que el valor del token coincide exactamente con el registrado en ms-auditoria (sin espacios ni saltos de línea).")
    numbered(doc, "Verifica que el header se llama exactamente X-App-Token (no X-App-Token-Service ni variantes).")
    numbered(doc, "Pide al administrador de ms-auditoria que confirme que tu token está registrado y activo.")
    numbered(doc, "Nunca incluyas el token con prefijo 'Bearer ' — el header X-App-Token va el token directamente.")

    heading2(doc, "Problema 2 — ms-auditoria devuelve 422 Unprocessable Entity")
    callout(doc, "Uno o más campos de tu log tienen formato inválido.", kind='error')
    numbered(doc, "Lee el campo data.invalid_fields en la respuesta 422 — te dice exactamente qué campos fallaron.")
    numbered(doc, "Verifica que timestamp está en formato ISO 8601 con zona UTC: 2026-03-02T14:05:12Z.")
    numbered(doc, "Verifica que response_code es un entero (no string) entre 100 y 599.")
    numbered(doc, "Verifica que duration_ms es un entero no negativo (no float, no string).")
    numbered(doc, "Verifica que no falten campos obligatorios (timestamp, service_name, functionality, method, response_code, duration_ms).")
    numbered(doc, "Recuerda que request_id y detail son opcionales en el contrato actual.")

    heading2(doc, "Problema 3 — El log se pierde (no aparece en ms-auditoria) pero no hay error")
    callout(doc, "El log se envió pero ocurrió un error al persistirlo en background.", kind='warning')
    numbered(doc, "Esto es comportamiento esperado por diseño: ms-auditoria puede perder logs en condiciones de fallo interno sin afectar a tu servicio.")
    numbered(doc, "Si necesitas garantía de entrega, implementa un buffer local con reintentos en tu microservicio.")
    numbered(doc, "Verifica los logs internos de ms-auditoria (stderr) para ver si hay errores de conexión a PostgreSQL.")

    heading2(doc, "Problema 4 — El envío del log está bloqueando la respuesta de tu endpoint")
    callout(doc, "Estás haciendo await del log sin usar BackgroundTasks o create_task.", kind='error')
    numbered(doc, "Busca en tu código si estás haciendo await send_audit_log(...) directamente en el handler.")
    numbered(doc, "Cambia a background_tasks.add_task(send_audit_log, ...) para envío en background.")
    numbered(doc, "Nunca uses await para llamadas a ms-auditoria dentro del handler principal del endpoint.")

    heading2(doc, "Problema 5 — ConnectionRefusedError o HTTPX Connect Error")
    callout(doc, "ms-auditoria no está disponible en la dirección configurada.", kind='warning')
    numbered(doc, "Verifica que AUD_BASE_URL apunta a la URL correcta (http://ms-auditoria:8019 en Docker Compose).")
    numbered(doc, "Si estás en desarrollo local, verifica que el contenedor de ms-auditoria está corriendo.")
    numbered(doc, "Haz un GET a http://ms-auditoria:8019/api/v1/health para confirmar si el servicio responde.")
    numbered(doc, "Tu servicio debe continuar operando normalmente aunque ms-auditoria no esté disponible.")

    heading2(doc, "Problema 6 — Los logs no tienen el Request ID correcto")
    callout(doc, "No estás propagando el X-Request-ID correctamente.", kind='warning')
    numbered(doc, "En tu audit_client.py, el Request ID debe venir de request.headers.get('x-request-id').")
    numbered(doc, "Si el header no existe en la petición entrante, genera uno nuevo (no uses un ID fijo).")
    numbered(doc, "El request_id del log debe ser el mismo que el X-Request-ID de la petición que llegó a TU endpoint.")


def build_sec15_checklist(doc):
    add_page_break(doc)
    heading1(doc, "15. Resumen de Reglas y Checklist de Integración")

    callout(doc,
        "Usa esta sección antes de hacer merge a main. Si todas las casillas están marcadas, "
        "la integración está correctamente implementada.",
        kind='success')

    heading2(doc, "15.1 Resumen de reglas obligatorias")
    info_table(doc,
        [
            ["Regla",        "Descripción",                                                      "Referencia"],
            ["RT-01",        "El envío de logs es asíncrono (fire-and-forget). Nunca bloquea.",  "Sección 3, Principio 1"],
            ["RT-02",        "Si ms-auditoria falla, tu servicio continúa sin errores.",         "Sección 3, Principio 2"],
            ["RT-03",        "Todas las operaciones generan log (no solo los errores).",         "Sección 3, Principio 3"],
            ["RT-04",        "El Request ID se propaga en todas las llamadas.",                  "Sección 3, Principio 4"],
            ["RT-05",        "El campo detail no contiene datos sensibles.",                     "Sección 3, Principio 5"],
            ["RT-06",        "El token de aplicación está en variable de entorno, no en código.","Sección 5"],
            ["RT-07",        "Timeout máximo de 5 segundos configurado para ms-auditoria.",      "Sección 4.3"],
            ["RT-08",        "El campo timestamp está en formato ISO 8601 UTC.",                 "Sección 6"],
            ["RT-09",        "Los campos response_code y duration_ms son enteros, no strings.",  "Sección 6"],
            ["RT-10",        "user_id es null para operaciones sin usuario, nunca string vacío.","Sección 6"],
        ],
        headers=["ID", "Regla", "Sección"],
        col_widths=[0.6, 4.5, 1.5],
    )

    heading2(doc, "15.2 Checklist de verificación")
    checks = [
        "[ ] Instalé httpx en requirements.txt de mi microservicio.",
        "[ ] Creé el módulo audit_client.py con la función send_audit_log.",
        "[ ] Configuré las variables de entorno: APP_TOKEN, SERVICE_NAME, AUD_BASE_URL.",
        "[ ] El APP_TOKEN NO está hardcodeado en el código fuente ni en archivos versionados.",
        "[ ] Agregué el envío del log como BackgroundTask (no con await directo) en todos mis endpoints.",
        "[ ] El log incluye todos los campos obligatorios del contrato actual.",
        "[ ] Si envío request_id, su longitud es ≤ 36 caracteres.",
        "[ ] Si envío detail, no contiene datos sensibles y no supera 5000 caracteres.",
        "[ ] El campo timestamp está en formato ISO 8601 UTC (termina en Z).",
        "[ ] El campo service_name tiene el nombre exacto de mi microservicio (ej: ms-reservas).",
        "[ ] El campo response_code es el mismo código HTTP que devolvió mi endpoint al cliente.",
        "[ ] El campo duration_ms mide el tiempo real desde que llegó la petición hasta que envié la respuesta.",
        "[ ] El campo user_id usa el ID del usuario de la sesión activa, o null si no aplica.",
        "[ ] El campo detail es descriptivo pero no contiene contraseñas, tokens ni datos sensibles.",
        "[ ] Probé que mi servicio funciona correctamente cuando ms-auditoria está caído.",
        "[ ] Probé que el Request ID se propaga correctamente entre microservicios.",
        "[ ] Revisé los logs de mi servicio y no hay errores de formato 422 al enviar logs.",
        "[ ] Verifiqué en ms-auditoria que los logs de mi microservicio aparecen correctamente.",
    ]
    for check in checks:
        bullet(doc, check)

    heading2(doc, "15.3 Contacto y soporte")
    body_text(doc,
        "Si tienes dudas sobre la integración, o si ms-auditoria devuelve un error que no "
        "puedes resolver con esta guía, contacta al equipo de infraestructura / arquitectura "
        "con la siguiente información:")
    bullet(doc, "Nombre de tu microservicio.")
    bullet(doc, "Descripción del problema y código de error recibido.")
    bullet(doc, "Payload del log que estás enviando (sin datos sensibles).")
    bullet(doc, "Request ID de la petición problemática.")

    add_horizontal_rule(doc)
    p = doc.add_paragraph()
    r = p.add_run("— Fin del documento —")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, 120, 60)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("ms-auditoria [AUD] — Guía de Integración v1.0 — Marzo 2026 — ERP Universitario")
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ════════════════════════════════════════════════════════════════════════════
#  MAIN
# ════════════════════════════════════════════════════════════════════════════

def main():
    doc = Document()

    # Márgenes del documento
    section = doc.sections[0]
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Fuente por defecto
    style = doc.styles['Normal']
    style.font.name  = 'Calibri'
    style.font.size  = Pt(11)
    style.font.color.rgb = C_BODY

    # Construir secciones
    build_cover(doc)
    build_toc(doc)
    build_sec1_intro(doc)
    build_sec2_que_es(doc)
    build_sec3_principios(doc)
    build_sec4_conexion(doc)
    build_sec5_auth(doc)
    build_sec6_estructura_log(doc)
    build_sec7_endpoint_individual(doc)
    build_sec8_endpoint_batch(doc)
    build_sec9_errores(doc)
    build_sec10_implementacion(doc)
    build_sec11_request_id(doc)
    build_sec12_health(doc)
    build_sec13_ejemplos_por_ms(doc)
    build_sec14_troubleshooting(doc)
    build_sec15_checklist(doc)

    doc.save(OUTPUT_PATH)
    print(f"✅ Documento generado: {OUTPUT_PATH}")
    import os
    size = os.path.getsize(OUTPUT_PATH)
    print(f"   Tamaño: {size:,} bytes ({size/1024:.1f} KB)")


if __name__ == "__main__":
    main()
