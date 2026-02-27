"""
Script para generar el Documento Técnico de Arquitectura en formato Word (.docx)
a partir del contenido del markdown, con formato profesional.
"""

import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Aplica bordes a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_formatted_text(paragraph, text):
    """Agrega texto con formato inline (negrita, código, cursiva)."""
    # Patrón para encontrar **bold**, `code`, *italic*
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            # Fondo gris claro para código inline - simulado con color
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

def create_document():
    doc = Document()
    
    # ── Configurar estilos ─────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    
    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ══════════════════════════════════════════════════════════════════════
    # PORTADA
    # ══════════════════════════════════════════════════════════════════════
    
    # Espaciado superior
    for _ in range(4):
        doc.add_paragraph()
    
    # Título principal
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Documento Técnico de Arquitectura')
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.bold = True
    
    # Subtítulo
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('ms-auditoria')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x2E, 0x54, 0x8C)
    run.bold = True
    
    # Línea separadora
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run('━' * 50)
    run.font.color.rgb = RGBColor(0x2E, 0x54, 0x8C)
    run.font.size = Pt(12)
    
    # Descripción
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('Microservicio #19: Auditoría y Logging del ERP Universitario')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    
    doc.add_paragraph()
    
    # Info de la portada
    info_lines = [
        ('Materia:', 'Desarrollo de Software 3'),
        ('Tecnología:', 'FastAPI + SQLAlchemy 2.0 Async + PostgreSQL 16'),
        ('Versión:', '1.0.0'),
        ('Fecha:', 'Febrero 2026'),
        ('Repositorio:', 'https://github.com/Karatsuyu/ms-auditoria'),
    ]
    
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label + ' ')
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run = p.add_run(value)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════════════
    # TABLA DE CONTENIDOS
    # ══════════════════════════════════════════════════════════════════════
    
    toc_title = doc.add_paragraph()
    run = toc_title.add_run('Tabla de Contenidos')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.bold = True
    
    doc.add_paragraph()
    
    toc_items = [
        '1. Información General',
        '2. Descripción General del Sistema',
        '3. Arquitectura de Alto Nivel',
        '4. Arquitectura Interna Detallada',
        '5. Modelo de Datos',
        '6. Seguridad',
        '7. Concurrencia y Rendimiento',
        '8. Testing y CI/CD',
        '9. DevOps y Despliegue',
        '10. Justificaciones Técnicas',
    ]
    
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2E, 0x54, 0x8C)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════════════
    # HELPER FUNCTIONS
    # ══════════════════════════════════════════════════════════════════════
    
    def add_section_title(text, level=1):
        """Agrega título de sección."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        if level == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            run.bold = True
            # Línea debajo
            sep = doc.add_paragraph()
            run2 = sep.add_run('─' * 70)
            run2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run2.font.size = Pt(8)
            sep.paragraph_format.space_after = Pt(8)
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x2E, 0x54, 0x8C)
            run.bold = True
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x44, 0x66, 0x99)
            run.bold = True
        return p
    
    def add_body(text):
        """Agrega párrafo de texto con formato inline."""
        p = doc.add_paragraph()
        add_formatted_text(p, text)
        return p
    
    def add_table(headers, rows, col_widths=None):
        """Agrega tabla formateada."""
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, '2E548C')
        
        # Rows
        for r_idx, row in enumerate(rows):
            bg = 'F5F7FA' if r_idx % 2 == 0 else 'FFFFFF'
            for c_idx, cell_text in enumerate(row):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                add_formatted_text(p, str(cell_text))
                for run in p.runs:
                    run.font.size = Pt(9)
                set_cell_shading(cell, bg)
        
        # Ancho de columnas
        if col_widths:
            for row in table.rows:
                for idx, width in enumerate(col_widths):
                    row.cells[idx].width = Inches(width)
        
        doc.add_paragraph()  # Espacio después
        return table
    
    def add_code_block(code, language=''):
        """Agrega bloque de código."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        return p
    
    def add_note(text):
        """Agrega nota/observación."""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run('📝 Nota: ')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x2E, 0x54, 0x8C)
        add_formatted_text(p, text)
        for run in p.runs[1:]:
            run.font.size = Pt(10)
        return p
    
    # ══════════════════════════════════════════════════════════════════════
    # SECCIÓN 1: INFORMACIÓN GENERAL
    # ══════════════════════════════════════════════════════════════════════
    
    add_section_title('1. Información General')
    
    add_table(
        ['Campo', 'Valor'],
        [
            ['Nombre', 'ms-auditoria'],
            ['Número', 'Microservicio #19'],
            ['Puerto', '8019'],
            ['Framework', 'FastAPI 0.115.6'],
            ['ORM', 'SQLAlchemy 2.0.36 (100% async)'],
            ['Base de datos', 'PostgreSQL 16'],
            ['Driver async', 'asyncpg 0.30.0'],
            ['Validación', 'Pydantic 2.10.3 + pydantic-settings 2.7.0'],
            ['Migraciones', 'Alembic 1.14.0'],
            ['Servidor ASGI', 'Uvicorn 0.34.0'],
            ['Python', '3.10'],
            ['Lenguaje', '100% Python, 100% async/await'],
        ],
        col_widths=[2.0, 4.5]
    )
    
    # ══════════════════════════════════════════════════════════════════════
    # SECCIÓN 2: DESCRIPCIÓN GENERAL DEL SISTEMA
    # ══════════════════════════════════════════════════════════════════════
    
    add_section_title('2. Descripción General del Sistema')
    
    add_section_title('2.1 Propósito', level=2)
    add_body('ms-auditoria es el microservicio centralizado de auditoría y logging del ERP Universitario. Es responsable de registrar, almacenar, consultar y analizar todos los eventos generados por los 18 microservicios restantes del sistema.')
    
    add_section_title('2.2 Responsabilidades', level=2)
    
    responsibilities = [
        ('Recepción de eventos', 'Recibe logs de auditoría vía HTTP REST desde cualquier microservicio autorizado.'),
        ('Persistencia', 'Almacena los eventos en PostgreSQL con modelo optimizado para consultas de alto volumen.'),
        ('Consulta', 'Proporciona endpoints con filtros avanzados, paginación y búsqueda full-text.'),
        ('Trazabilidad', 'Permite rastrear una petición a través de múltiples microservicios usando X-Request-ID.'),
        ('Estadísticas', 'Genera métricas (logs por servicio, tasa de errores, duración promedio).'),
        ('Retención automática (TTL)', 'Purga automática de logs antiguos configurable.'),
        ('Seguridad', 'Autenticación inter-servicio con API Keys, cifrado AES-256-GCM disponible, rate limiting por IP.'),
    ]
    for title, desc in responsibilities:
        p = doc.add_paragraph()
        run = p.add_run(f'• {title}: ')
        run.bold = True
        p.add_run(desc)
    
    add_section_title('2.3 Posición en el ERP', level=2)
    add_body('Todos los 18 microservicios del ERP envían sus eventos de auditoría a ms-auditoria vía POST /api/v1/audit/log, autenticándose con un API Key en el header X-API-Key. ms-auditoria centraliza estos logs en PostgreSQL para consulta, trazabilidad y análisis.')
    
    add_code_block(
        '┌──────────────────────────────────────────────────────────────┐\n'
        '│                     ERP UNIVERSITARIO                        │\n'
        '│                                                              │\n'
        '│  ms-autenticación ──┐                                        │\n'
        '│  ms-roles ──────────┤                                        │\n'
        '│  ms-usuarios ───────┤                                        │\n'
        '│  ms-académica ──────┤     POST /api/v1/audit/log             │\n'
        '│  ms-matrículas ─────┤────────────────────────►┌────────────┐ │\n'
        '│  ms-calificaciones ─┤                         │ms-auditoria│ │\n'
        '│  ms-horarios ───────┤  (X-API-Key + JSON)     │   :8019    │ │\n'
        '│  ms-pagos ──────────┤                         │            │ │\n'
        '│  ms-becas ──────────┤                         │ PostgreSQL │ │\n'
        '│  ms-biblioteca ─────┤                         │   :5432    │ │\n'
        '│  ... (18 total) ────┘                         └────────────┘ │\n'
        '└──────────────────────────────────────────────────────────────┘'
    )
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════════════
    # SECCIÓN 3: ARQUITECTURA DE ALTO NIVEL
    # ══════════════════════════════════════════════════════════════════════
    
    add_section_title('3. Arquitectura de Alto Nivel')
    
    add_section_title('3.1 Estilo Arquitectónico', level=2)
    add_body('El microservicio implementa una **arquitectura en capas (Clean Architecture)** con 4 capas bien definidas:')
    
    add_code_block(
        '┌─────────────────────────────────────┐\n'
        '│         ROUTES (Presentación)       │  ← audit_routes.py\n'
        '│   FastAPI Router + Endpoints        │\n'
        '├─────────────────────────────────────┤\n'
        '│        SERVICES (Negocio)           │  ← audit_service.py, statistics_service.py\n'
        '│   Lógica de negocio + orquestación  │\n'
        '├─────────────────────────────────────┤\n'
        '│      REPOSITORIES (Datos)           │  ← audit_repository.py\n'
        '│   Consultas SQL async (select/func) │\n'
        '├─────────────────────────────────────┤\n'
        '│         MODELS (Dominio)            │  ← audit_log.py, microservice_token.py\n'
        '│   ORM SQLAlchemy + Schemas Pydantic │\n'
        '└─────────────────────────────────────┘'
    )
    
    add_section_title('3.2 Patrones de Diseño Implementados', level=2)
    
    add_table(
        ['Patrón', 'Implementación', 'Archivo(s)'],
        [
            ['Repository Pattern', 'AuditRepository abstrae todas las consultas SQL', 'repositories/audit_repository.py'],
            ['Unit of Work', 'Disponible como infraestructura para transacciones atómicas complejas; endpoints actuales usan AsyncSession directa', 'database/unit_of_work.py'],
            ['Dependency Injection', 'FastAPI Depends() para sesiones de BD y autenticación', 'core/dependencies.py'],
            ['CQRS-like', 'Separación de comandos (POST/DELETE) y queries (GET)', 'routes/audit_routes.py'],
            ['Middleware Chain', 'Cadena de middlewares para cross-cutting concerns', 'core/middleware.py, core/rate_limiter.py'],
            ['Strategy Pattern', 'GUID TypeDecorator adapta UUID entre PostgreSQL y SQLite', 'models/audit_log.py'],
            ['Singleton', 'Instancias únicas de retention_service, aes_cipher, logger, settings', 'Varios módulos'],
        ],
        col_widths=[1.5, 3.0, 2.0]
    )
    
    add_section_title('3.3 Estructura del Proyecto', level=2)
    
    add_code_block(
        'ms-auditoria/\n'
        '├── app/\n'
        '│   ├── core/                          # Configuración y cross-cutting concerns\n'
        '│   │   ├── config.py                  # Pydantic Settings (multi-entorno)\n'
        '│   │   ├── middleware.py              # X-Request-ID + response time\n'
        '│   │   ├── rate_limiter.py            # Rate limiting por IP (sliding window)\n'
        '│   │   ├── auth.py                    # Autenticación inter-servicio (API Keys)\n'
        '│   │   ├── security.py               # AES-256-GCM cifrado\n'
        '│   │   ├── dependencies.py           # FastAPI Dependency Injection\n'
        '│   │   └── exception_handlers.py     # Manejadores globales de errores\n'
        '│   ├── database/                      # Capa de acceso a datos\n'
        '│   │   ├── base.py                    # SQLAlchemy DeclarativeBase\n'
        '│   │   ├── connection.py              # AsyncEngine + SyncEngine\n'
        '│   │   ├── session.py                 # AsyncSessionLocal factory\n'
        '│   │   └── unit_of_work.py           # Patrón Unit of Work\n'
        '│   ├── models/                        # Modelos ORM\n'
        '│   │   ├── audit_log.py              # Tabla audit_logs (GUID cross-DB)\n'
        '│   │   └── microservice_token.py     # Tabla microservice_tokens\n'
        '│   ├── repositories/                  # Repository Pattern (async)\n'
        '│   │   └── audit_repository.py\n'
        '│   ├── schemas/                       # Pydantic v2 schemas\n'
        '│   │   ├── audit_schema.py           # Create, Response, Filter\n'
        '│   │   └── response_schema.py        # Respuestas genéricas\n'
        '│   ├── services/                      # Lógica de negocio\n'
        '│   │   ├── audit_service.py          # Servicio principal\n'
        '│   │   ├── statistics_service.py     # Estadísticas y métricas\n'
        '│   │   ├── auth_service.py           # Comunicación con ms-autenticación\n'
        '│   │   └── retention_service.py      # TTL / purga automática\n'
        '│   ├── routes/\n'
        '│   │   └── audit_routes.py           # 9 endpoints RESTful\n'
        '│   ├── utils/\n'
        '│   │   └── logger.py                 # JSON structured logging\n'
        '│   └── main.py                        # Punto de entrada FastAPI\n'
        '├── alembic/                           # Migraciones de BD\n'
        '├── tests/                             # 38 unit + 12 integration tests\n'
        '├── .github/workflows/ci.yml          # CI/CD pipeline (4 jobs)\n'
        '├── Dockerfile                         # Multi-stage build\n'
        '├── docker-compose.yml                 # PostgreSQL 16 + App\n'
        '└── requirements.txt'
    )
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════════════
    # SECCIÓN 4: ARQUITECTURA INTERNA DETALLADA
    # ══════════════════════════════════════════════════════════════════════
    
    add_section_title('4. Arquitectura Interna Detallada')
    
    add_section_title('4.1 Flujo de una Petición (POST /api/v1/audit/log)', level=2)
    
    add_code_block(
        'Cliente (ms-matriculas)\n'
        '    │\n'
        '    │  POST /api/v1/audit/log\n'
        '    │  Headers: X-API-Key: <token>, X-Request-ID: <uuid>\n'
        '    │  Body: { timestamp, nombre_microservicio, endpoint, ... }\n'
        '    │\n'
        '    ▼\n'
        '┌─────────────────────────────────────┐\n'
        '│  1. RequestIDMiddleware             │  Inyecta/propaga X-Request-ID\n'
        '│     (core/middleware.py)            │  Mide tiempo con perf_counter\n'
        '├─────────────────────────────────────┤\n'
        '│  2. RateLimitMiddleware             │  Sliding window por IP\n'
        '│     (core/rate_limiter.py)          │  429 si excede límite\n'
        '├─────────────────────────────────────┤\n'
        '│  3. CORSMiddleware                  │  Valida origen permitido\n'
        '│     (Starlette built-in)            │\n'
        '├─────────────────────────────────────┤\n'
        '│  4. ExceptionHandlers               │  Captura errores globales\n'
        '│     (core/exception_handlers.py)    │\n'
        '├─────────────────────────────────────┤\n'
        '│  5. Route: create_audit_log()       │  Endpoint FastAPI\n'
        '│     - Depends(get_db) → AsyncSession│\n'
        '│     - Depends(verify_api_key) → auth│\n'
        '├─────────────────────────────────────┤\n'
        '│  6. AuditService.create_log()       │  Lógica de negocio\n'
        '│     - detalle se persiste como texto│\n'
        '├─────────────────────────────────────┤\n'
        '│  7. AuditRepository.save()          │  Persistencia async\n'
        '│     - session.add → flush → refresh │\n'
        '│     - await db.commit()             │\n'
        '├─────────────────────────────────────┤\n'
        '│  8. PostgreSQL 16                   │  INSERT en audit_logs\n'
        '│     (asyncpg driver)                │  8 índices optimizados\n'
        '└─────────────────────────────────────┘\n'
        '    │\n'
        '    ▼\n'
        'Response 201 Created\n'
        'Headers: X-Request-ID, X-Response-Time-ms, X-RateLimit-*\n'
        'Body: { success: true, data: { id, servicio, ... } }'
    )
    
    add_section_title('4.2 Endpoints Implementados (9 totales)', level=2)
    
    add_table(
        ['#', 'Método', 'Ruta', 'Función', 'Auth', 'Descripción'],
        [
            ['1', 'GET', '/api/v1/audit/health', 'health_check', 'No', 'Health check'],
            ['2', 'POST', '/api/v1/audit/log', 'create_audit_log', 'API Key', 'Registrar evento'],
            ['3', 'POST', '/api/v1/audit/log/batch', 'create_audit_logs_batch', 'API Key', 'Registro masivo (máx 1000)'],
            ['4', 'GET', '/api/v1/audit/logs', 'get_audit_logs', 'No', 'Listar con filtros + paginación'],
            ['5', 'GET', '/api/v1/audit/log/{audit_id}', 'get_audit_log_by_id', 'No', 'Obtener log por UUID'],
            ['6', 'GET', '/api/v1/audit/trace/{request_id}', 'trace_request', 'No', 'Trazabilidad por X-Request-ID'],
            ['7', 'GET', '/api/v1/audit/user/{usuario_id}', 'get_user_audit_logs', 'No', 'Historial de usuario'],
            ['8', 'GET', '/api/v1/audit/stats', 'get_statistics', 'No', 'Estadísticas generales'],
            ['9', 'DELETE', '/api/v1/audit/purge', 'purge_logs', 'API Key', 'Purgar logs antiguos'],
        ],
        col_widths=[0.3, 0.5, 1.8, 1.5, 0.6, 1.5]
    )
    
    add_body('Adicionalmente, existe un endpoint raíz GET / definido en main.py con información del microservicio.')
    
    add_section_title('4.3 Ejemplo del Endpoint Principal', level=2)
    
    add_code_block(
        '# routes/audit_routes.py — POST /api/v1/audit/log\n\n'
        '@router.post(\n'
        '    "/log",\n'
        '    response_model=DataResponse[AuditLogResponse],\n'
        '    status_code=status.HTTP_201_CREATED,\n'
        '    summary="Registrar evento de auditoría",\n'
        ')\n'
        'async def create_audit_log(\n'
        '    data: AuditLogCreate,\n'
        '    db: AsyncSession = Depends(get_db),\n'
        '    _api_key=Depends(verify_api_key),\n'
        '):\n'
        '    service = AuditService(db)\n'
        '    result = await service.create_log(data)\n'
        '    return DataResponse(\n'
        '        success=True,\n'
        '        message="Evento de auditoría registrado exitosamente",\n'
        '        data=result,\n'
        '    )'
    )
    
    add_body('**Dependencias inyectadas:**')
    p = doc.add_paragraph()
    run = p.add_run('• get_db: ')
    run.bold = True
    p.add_run('Provee una AsyncSession por request, se cierra automáticamente al finalizar.')
    
    p = doc.add_paragraph()
    run = p.add_run('• verify_api_key: ')
    run.bold = True
    p.add_run('Valida el header X-API-Key contra la tabla microservice_tokens usando SHA-256. En development/testing permite acceso sin key.')
    
    add_section_title('4.4 Capa de Servicio — Mapeo Schema → ORM', level=2)
    
    add_code_block(
        '# services/audit_service.py\n\n'
        'async def create_log(self, data: AuditLogCreate) -> AuditLogResponse:\n'
        '    audit_log = AuditLog(\n'
        '        request_id=data.request_id or str(uuid.uuid4()),\n'
        '        servicio=data.nombre_microservicio,\n'
        '        endpoint=data.endpoint,\n'
        '        metodo=data.metodo_http,\n'
        '        codigo_respuesta=data.codigo_respuesta,\n'
        '        duracion_ms=data.duracion_ms,\n'
        '        usuario_id=data.usuario_id,\n'
        '        detalle=data.detalle,\n'
        '        ip_origen=data.ip_origen,\n'
        '        timestamp_evento=data.timestamp,\n'
        '    )\n'
        '    saved = await self.repo.save(audit_log)\n'
        '    await self.db.commit()\n'
        '    return AuditLogResponse.model_validate(saved)'
    )
    
    add_note('El campo detalle se persiste como texto plano para permitir la búsqueda full-text con el índice GIN de PostgreSQL. El módulo core/security.py provee la clase AESCipher con cifrado AES-256-GCM disponible como utilidad si se requiere cifrar datos especialmente sensibles en el futuro, pero no se aplica automáticamente en el flujo de creación.')
    
    add_section_title('4.5 Cadena de Middleware', level=2)
    
    add_body('Los middlewares se registran en main.py. Starlette ejecuta los middlewares en **orden inverso** al registro (LIFO). El orden de ejecución real para una petición entrante es:')
    
    add_code_block(
        'Request entrante\n'
        '    → RequestIDMiddleware  (inyecta X-Request-ID, mide tiempo)\n'
        '    → RateLimitMiddleware  (sliding window por IP, 429 si excede)\n'
        '    → CORSMiddleware       (valida origen, agrega headers CORS)\n'
        '    → Endpoint\n'
        '    → (respuesta sube por la misma cadena en orden inverso)'
    )
    
    add_section_title('4.6 Exception Handlers Globales', level=2)
    
    add_body('Registrados en core/exception_handlers.py mediante `register_exception_handlers(app)`:')
    
    add_table(
        ['Handler', 'Captura', 'Respuesta'],
        [
            ['http_exception_handler', 'StarletteHTTPException (4xx, 5xx)', 'JSON con success: false'],
            ['validation_exception_handler', 'RequestValidationError (Pydantic)', '422 con lista de errores por campo'],
            ['unhandled_exception_handler', 'Exception genérica', '500 con detalle oculto en producción'],
        ],
        col_widths=[2.0, 2.2, 2.3]
    )
    
    add_body('Formato de respuesta de error estandarizado:')
    add_code_block(
        '{\n'
        '  "success": false,\n'
        '  "error": "Descripción del tipo de error",\n'
        '  "detail": "Detalle específico del error"\n'
        '}'
    )
    
    add_section_title('4.7 Servicio de Estadísticas', level=2)
    
    add_body('`StatisticsService` genera métricas consultando el repositorio:')
    
    add_table(
        ['Métrica', 'Método del repositorio', 'Descripción'],
        [
            ['Total de registros', 'count_total()', 'COUNT total de audit_logs'],
            ['Logs por servicio', 'count_by_servicio()', 'GROUP BY servicio, ORDER BY total DESC'],
            ['Logs por código HTTP', 'count_by_codigo_respuesta()', 'GROUP BY codigo_respuesta'],
            ['Duración promedio', 'average_duration_by_servicio()', 'AVG(duracion_ms) por servicio'],
            ['Tasa de errores', 'error_rate_by_servicio()', '% de códigos ≥400 por servicio'],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )
    
    add_section_title('4.8 Servicio de Retención Automática (TTL)', level=2)
    
    retention_points = [
        ('Scheduler', 'Loop con asyncio.create_task() — sin dependencias externas (no usa APScheduler ni Celery).'),
        ('Ejecución', 'Calcula segundos hasta RETENTION_CRON_HOUR (default: 03:00 UTC), duerme con asyncio.sleep(), ejecuta purga.'),
        ('Purga', 'DELETE FROM audit_logs WHERE timestamp_evento < (now - RETENTION_DAYS).'),
        ('Sesión propia', 'Usa AsyncSessionLocal() independiente de los requests HTTP.'),
        ('Lifecycle', 'Se inicia en lifespan startup, se detiene en shutdown con task.cancel().'),
        ('Resiliencia', 'En caso de error, espera 1 hora antes de reintentar.'),
    ]
    for title_text, desc_text in retention_points:
        p = doc.add_paragraph()
        run = p.add_run(f'• {title_text}: ')
        run.bold = True
        p.add_run(desc_text)
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════════════
    # SECCIÓN 5: MODELO DE DATOS
    # ══════════════════════════════════════════════════════════════════════
    
    add_section_title('5. Modelo de Datos')
    
    add_section_title('5.1 Diagrama Entidad-Relación', level=2)
    
    add_code_block(
        '┌─────────────────────────────────────────────────┐\n'
        '│                  audit_logs                      │\n'
        '├─────────────────────────────────────────────────┤\n'
        '│ PK  id                UUID          NOT NULL     │\n'
        '│     request_id        VARCHAR(50)   NOT NULL     │\n'
        '│     servicio          VARCHAR(50)   NOT NULL     │\n'
        '│     endpoint          VARCHAR(200)  NOT NULL     │\n'
        '│     metodo            VARCHAR(10)   NOT NULL     │\n'
        '│     codigo_respuesta  INTEGER       NULLABLE     │\n'
        '│     duracion_ms       INTEGER       NULLABLE     │\n'
        '│     usuario_id        UUID          NULLABLE     │\n'
        '│     detalle           TEXT          NULLABLE     │\n'
        '│     ip_origen         VARCHAR(45)   NULLABLE     │\n'
        '│     timestamp_evento  TIMESTAMP(tz) NOT NULL     │\n'
        '│     created_at        TIMESTAMP(tz) NOT NULL     │\n'
        '├─────────────────────────────────────────────────┤\n'
        '│ Índices: 8 (4 simples + 3 compuestos + 1 GIN)  │\n'
        '└─────────────────────────────────────────────────┘\n'
        '\n'
        '┌─────────────────────────────────────────────────┐\n'
        '│              microservice_tokens                 │\n'
        '├─────────────────────────────────────────────────┤\n'
        '│ PK  id                    UUID          NOT NULL │\n'
        '│ UQ  nombre_microservicio  VARCHAR(50)   NOT NULL │\n'
        '│     token_hash            VARCHAR(256)  NOT NULL │\n'
        '│     activo                BOOLEAN       NOT NULL │\n'
        '│     created_at            TIMESTAMP(tz) NOT NULL │\n'
        '│     updated_at            TIMESTAMP(tz) NOT NULL │\n'
        '└─────────────────────────────────────────────────┘'
    )
    
    add_section_title('5.2 Tabla audit_logs — Detalle de Columnas', level=2)
    
    add_table(
        ['Columna', 'Tipo SQL', 'Tipo ORM', 'Nullable', 'Default', 'Comentario'],
        [
            ['id', 'UUID / CHAR(36)', 'GUID() custom', 'NOT NULL', 'uuid.uuid4()', 'Identificador único'],
            ['request_id', 'VARCHAR(50)', 'String(50)', 'NOT NULL', '—', 'X-Request-ID trazabilidad'],
            ['servicio', 'VARCHAR(50)', 'String(50)', 'NOT NULL', '—', 'Microservicio emisor'],
            ['endpoint', 'VARCHAR(200)', 'String(200)', 'NOT NULL', '—', 'Ruta del endpoint'],
            ['metodo', 'VARCHAR(10)', 'String(10)', 'NOT NULL', '—', 'Método HTTP'],
            ['codigo_respuesta', 'INTEGER', 'Integer', 'NULLABLE', '—', 'Código HTTP'],
            ['duracion_ms', 'INTEGER', 'Integer', 'NULLABLE', '—', 'Duración en ms'],
            ['usuario_id', 'UUID / CHAR(36)', 'GUID()', 'NULLABLE', 'None', 'UUID del usuario'],
            ['detalle', 'TEXT', 'Text', 'NULLABLE', 'None', 'Detalle adicional (JSON)'],
            ['ip_origen', 'VARCHAR(45)', 'String(45)', 'NULLABLE', 'None', 'IP de origen (IPv4/IPv6)'],
            ['timestamp_evento', 'TIMESTAMP(tz)', 'TIMESTAMP(tz=True)', 'NOT NULL', '—', 'Momento del evento'],
            ['created_at', 'TIMESTAMP(tz)', 'TIMESTAMP(tz=True)', 'NOT NULL', 'datetime.now(UTC)', 'Momento de registro'],
        ],
        col_widths=[1.1, 1.0, 1.0, 0.7, 1.0, 1.3]
    )
    
    add_note('El tipo GUID es un TypeDecorator custom que usa UUID nativo en PostgreSQL y CHAR(36) en SQLite, permitiendo que los unit tests funcionen con SQLite sin cambiar el modelo.')
    
    add_section_title('5.3 Tabla microservice_tokens — Detalle de Columnas', level=2)
    
    add_table(
        ['Columna', 'Tipo SQL', 'Tipo ORM', 'Nullable', 'Default', 'Comentario'],
        [
            ['id', 'UUID / CHAR(36)', 'GUID()', 'NOT NULL', 'uuid.uuid4()', 'ID del token'],
            ['nombre_microservicio', 'VARCHAR(50) UQ', 'String(50)', 'NOT NULL', '—', 'Nombre del microservicio'],
            ['token_hash', 'VARCHAR(256)', 'String(256)', 'NOT NULL', '—', 'Hash SHA-256 del API Key'],
            ['activo', 'BOOLEAN', 'Boolean', 'NOT NULL', 'True', 'Si está autorizado'],
            ['created_at', 'TIMESTAMP(tz)', 'TIMESTAMP(tz=True)', 'NOT NULL', 'datetime.now(UTC)', 'Fecha de creación'],
            ['updated_at', 'TIMESTAMP(tz)', 'TIMESTAMP(tz=True)', 'NOT NULL', 'datetime.now(UTC)', 'Última actualización (onupdate)'],
        ],
        col_widths=[1.3, 1.0, 0.9, 0.7, 1.0, 1.3]
    )
    
    add_section_title('5.4 Índices (8 totales)', level=2)
    
    add_section_title('Índices simples (4):', level=3)
    add_table(
        ['Nombre', 'Columna', 'Propósito'],
        [
            ['ix_audit_logs_request_id', 'request_id', 'Búsqueda por X-Request-ID (trazabilidad)'],
            ['ix_audit_logs_servicio', 'servicio', 'Filtro por microservicio emisor'],
            ['ix_audit_logs_codigo_respuesta', 'codigo_respuesta', 'Filtro por código HTTP'],
            ['ix_audit_logs_usuario_id', 'usuario_id', 'Filtro por usuario'],
        ],
        col_widths=[2.2, 1.3, 3.0]
    )
    
    add_section_title('Índices compuestos (3):', level=3)
    add_table(
        ['Nombre', 'Columnas', 'Propósito'],
        [
            ['ix_audit_servicio_timestamp', 'servicio, timestamp_evento', 'Logs por servicio en rango de tiempo'],
            ['ix_audit_usuario_timestamp', 'usuario_id, timestamp_evento', 'Historial de usuario en rango de tiempo'],
            ['ix_audit_codigo_servicio', 'codigo_respuesta, servicio', 'Estadísticas de errores por servicio'],
        ],
        col_widths=[2.2, 2.0, 2.3]
    )
    
    add_section_title('Índice GIN (1):', level=3)
    add_table(
        ['Nombre', 'Expresión', 'Propósito'],
        [
            ['ix_audit_detalle_fulltext', "GIN(to_tsvector('spanish', COALESCE(detalle, '')))", 'Búsqueda full-text en español'],
        ],
        col_widths=[2.2, 2.5, 1.8]
    )
    
    add_note('El índice GIN se crea en la migración Alembic b2a3c4d5e6f7 usando op.execute() con SQL raw porque es un índice funcional de PostgreSQL no soportado por el autogenerate de Alembic.')
    
    add_section_title('5.5 Migraciones Alembic', level=2)
    
    add_table(
        ['Revisión', 'ID', 'Descripción'],
        [
            ['1', 'fae4016df4b8', 'Schema inicial: tablas audit_logs y microservice_tokens con todos los índices'],
            ['2', 'b2a3c4d5e6f7', 'Índice GIN full-text en campo detalle para búsqueda en español'],
        ],
        col_widths=[0.8, 1.5, 4.2]
    )
    
    add_section_title('5.6 Schemas Pydantic', level=2)
    
    add_section_title('AuditLogCreate (entrada — POST)', level=3)
    
    add_table(
        ['Campo', 'Tipo', 'Requerido', 'Validación', 'Mapea a columna ORM'],
        [
            ['timestamp', 'datetime', 'Sí', 'ISO 8601', 'timestamp_evento'],
            ['nombre_microservicio', 'str', 'Sí', '1-50 chars', 'servicio'],
            ['endpoint', 'str', 'Sí', '1-200 chars', 'endpoint'],
            ['metodo_http', 'str', 'Sí', '1-10 chars', 'metodo'],
            ['codigo_respuesta', 'int', 'Sí', '100-599', 'codigo_respuesta'],
            ['duracion_ms', 'int', 'Sí', '≥0', 'duracion_ms'],
            ['usuario_id', 'UUID?', 'No', 'UUID válido', 'usuario_id'],
            ['detalle', 'str?', 'No', 'máx 5000 chars', 'detalle'],
            ['ip_origen', 'str?', 'No', 'máx 45 chars', 'ip_origen'],
            ['request_id', 'str?', 'No', 'máx 50 chars', 'request_id (auto UUID si null)'],
        ],
        col_widths=[1.3, 0.6, 0.6, 1.0, 1.8]
    )
    
    add_note('Los nombres de los campos del schema difieren de los nombres de las columnas ORM. El mapeo se realiza explícitamente en AuditService.create_log().')
    
    add_section_title('AuditLogResponse (salida — GET)', level=3)
    
    add_table(
        ['Campo', 'Tipo', 'Descripción'],
        [
            ['id', 'UUID', 'ID del registro'],
            ['request_id', 'str', 'X-Request-ID'],
            ['servicio', 'str', 'Microservicio emisor'],
            ['endpoint', 'str', 'Ruta invocada'],
            ['metodo', 'str', 'Método HTTP'],
            ['codigo_respuesta', 'int', 'Código HTTP'],
            ['duracion_ms', 'int', 'Duración en ms'],
            ['usuario_id', 'UUID?', 'UUID del usuario'],
            ['detalle', 'str?', 'Detalle del evento'],
            ['ip_origen', 'str?', 'IP de origen'],
            ['timestamp_evento', 'datetime', 'Momento del evento'],
            ['created_at', 'datetime', 'Momento de registro'],
        ],
        col_widths=[1.3, 1.0, 4.2]
    )
    
    add_body('Usa `ConfigDict(from_attributes=True)` para mapear directamente desde el modelo ORM.')
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════════════
    # SECCIÓN 6: SEGURIDAD
    # ══════════════════════════════════════════════════════════════════════
    
    add_section_title('6. Seguridad')
    
    add_section_title('6.1 Autenticación Inter-Servicio (API Keys)', level=2)
    add_body('**Archivo:** `core/auth.py`')
    
    add_table(
        ['Aspecto', 'Detalle'],
        [
            ['Mecanismo', 'API Key enviada en header X-API-Key'],
            ['Hash', 'SHA-256 (hashlib.sha256) del API Key'],
            ['Almacenamiento', 'Tabla microservice_tokens — solo se guarda el hash'],
            ['Validación', 'Busca token activo (activo=True) cuyo token_hash coincida'],
            ['Modo desarrollo', 'Si APP_ENV es development/testing y no hay key, permite acceso'],
            ['Modo producción', 'API Key obligatorio — retorna 401 si falta o es inválido'],
            ['Endpoints protegidos', 'POST /log, POST /log/batch, DELETE /purge'],
            ['Endpoints públicos', 'GET (health, logs, trace, user, stats)'],
        ],
        col_widths=[1.8, 4.7]
    )
    
    add_code_block(
        'X-API-Key header → SHA-256 hash → SELECT FROM microservice_tokens\n'
        '                                   WHERE token_hash = hash AND activo = True\n'
        '                                   → 200 OK / 401 Unauthorized'
    )
    
    add_section_title('6.2 Cifrado AES-256-GCM (Disponible)', level=2)
    add_body('**Archivo:** `core/security.py`')
    
    add_table(
        ['Aspecto', 'Detalle'],
        [
            ['Algoritmo', 'AES-256-GCM (cifrado autenticado)'],
            ['Librería', 'cryptography.hazmat.primitives.ciphers.aead.AESGCM'],
            ['Clave', 'AES_SECRET_KEY — 64 caracteres hexadecimales (256 bits)'],
            ['Nonce', '12 bytes random (os.urandom(12))'],
            ['Formato', 'Base64(nonce[12] + ciphertext)'],
            ['Instancia', 'Singleton aes_cipher disponible para importar'],
            ['Uso actual', 'Disponible como utilidad; no se aplica automáticamente. El campo detalle se persiste como texto plano para permitir búsqueda full-text con el índice GIN'],
        ],
        col_widths=[1.5, 5.0]
    )
    
    add_section_title('6.3 CORS (Cross-Origin Resource Sharing)', level=2)
    
    add_table(
        ['Aspecto', 'Detalle'],
        [
            ['Orígenes', 'Configurados vía CORS_ORIGINS (default: localhost:3000, localhost:8080)'],
            ['Desarrollo', 'allow_origins=["*"] y allow_credentials=False'],
            ['Producción', 'Orígenes específicos con allow_credentials=True'],
            ['Métodos', 'GET, POST, PUT, DELETE, PATCH, OPTIONS'],
            ['Headers expuestos', 'X-Request-ID, X-Response-Time-ms, X-RateLimit-Limit, X-RateLimit-Remaining'],
        ],
        col_widths=[1.5, 5.0]
    )
    
    add_section_title('6.4 Rate Limiting', level=2)
    add_body('**Archivo:** `core/rate_limiter.py`')
    
    add_table(
        ['Aspecto', 'Detalle'],
        [
            ['Algoritmo', 'Sliding window por IP'],
            ['Almacenamiento', 'En memoria (diccionario IP → [timestamps])'],
            ['Límite default', '100 requests / 60 segundos (configurable)'],
            ['IP real', 'Soporta X-Forwarded-For para proxies'],
            ['Excluidos', '/api/v1/audit/health, /docs, /redoc, /openapi.json, /'],
            ['Respuesta 429', 'JSON con Retry-After, X-RateLimit-Limit, X-RateLimit-Remaining'],
        ],
        col_widths=[1.5, 5.0]
    )
    
    add_section_title('6.5 X-Request-ID (Trazabilidad)', level=2)
    add_body('**Archivo:** `core/middleware.py`')
    
    traceability_points = [
        'Si el cliente envía X-Request-ID, se propaga.',
        'Si no lo envía, se genera automáticamente con uuid.uuid4().',
        'Se inyecta en request.state.request_id para uso dentro del request.',
        'Se retorna en los headers X-Request-ID y X-Response-Time-ms.',
    ]
    for point in traceability_points:
        p = doc.add_paragraph()
        p.add_run(f'• {point}')
    
    add_section_title('6.6 Docker — Ejecución No-Root', level=2)
    add_body('El Dockerfile crea un usuario appuser en grupo appgroup y ejecuta la aplicación como usuario no-root:')
    add_code_block(
        'RUN addgroup --system appgroup && adduser --system --ingroup appgroup appuser\n'
        'USER appuser'
    )
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════════════
    # SECCIÓN 7: CONCURRENCIA Y RENDIMIENTO
    # ══════════════════════════════════════════════════════════════════════
    
    add_section_title('7. Concurrencia y Rendimiento')
    
    add_section_title('7.1 Motor Async', level=2)
    
    add_table(
        ['Componente', 'Implementación'],
        [
            ['AsyncEngine', 'create_async_engine() de SQLAlchemy 2.0'],
            ['Driver', 'asyncpg — driver PostgreSQL nativo async'],
            ['Session Factory', 'async_sessionmaker(bind=async_engine, class_=AsyncSession)'],
            ['Opciones de sesión', 'autoflush=False, autocommit=False, expire_on_commit=False'],
        ],
        col_widths=[1.8, 4.7]
    )
    
    add_section_title('7.2 Pool de Conexiones', level=2)
    add_body('**Archivo:** `database/connection.py`')
    
    add_table(
        ['Parámetro', 'Default', 'Variable de entorno', 'Descripción'],
        [
            ['pool_size', '10', 'DB_POOL_SIZE', 'Conexiones activas en el pool'],
            ['max_overflow', '20', 'DB_MAX_OVERFLOW', 'Conexiones extra bajo alta carga'],
            ['pool_recycle', '3600', 'DB_POOL_RECYCLE', 'Reciclar conexiones cada N segundos'],
            ['pool_pre_ping', 'True', '—', 'Verificar conexión antes de usarla'],
        ],
        col_widths=[1.3, 0.8, 1.8, 2.6]
    )
    
    add_note('Para SQLite (usado en tests), se usa StaticPool y check_same_thread=False en vez del pool estándar.')
    
    add_section_title('7.3 Compatibilidad PostgreSQL / SQLite', level=2)
    add_body('El sistema detecta automáticamente el tipo de base de datos en `database/connection.py`:')
    
    add_code_block(
        '_is_sqlite = settings.DATABASE_URL.startswith("sqlite")\n\n'
        'if _is_sqlite:\n'
        '    async_engine = create_async_engine(..., poolclass=StaticPool)\n'
        'else:\n'
        '    async_engine = create_async_engine(..., pool_size=..., max_overflow=...)'
    )
    
    add_body('Esto permite que los **unit tests** usen SQLite en memoria y la **aplicación real** use PostgreSQL con pool optimizado.')
    
    add_section_title('7.4 Conversión Automática de URL', level=2)
    add_body('**Archivo:** `core/config.py` — `computed_field`')
    add_body('Solo se configura DATABASE_URL (sync). La URL async se genera automáticamente reemplazando el driver:')
    add_code_block(
        'postgresql+psycopg2://... → postgresql+asyncpg://...\n'
        'sqlite:///...             → sqlite+aiosqlite:///...'
    )
    
    add_section_title('7.5 Uvicorn en Producción', level=2)
    
    add_table(
        ['Opción', 'Valor', 'Propósito'],
        [
            ['--workers', '4', 'Procesos worker para paralelismo real'],
            ['--loop', 'uvloop', 'Event loop optimizado (más rápido que asyncio default)'],
            ['--http', 'httptools', 'Parser HTTP en C (más rápido que h11)'],
        ],
        col_widths=[1.3, 1.2, 4.0]
    )
    
    add_section_title('7.6 Resource Limits (Docker Compose)', level=2)
    add_code_block(
        'deploy:\n'
        '  resources:\n'
        '    limits:\n'
        '      cpus: "1.0"\n'
        '      memory: 512M'
    )
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════════════
    # SECCIÓN 8: TESTING Y CI/CD
    # ══════════════════════════════════════════════════════════════════════
    
    add_section_title('8. Testing y CI/CD')
    
    add_section_title('8.1 Estrategia de Testing', level=2)
    
    add_table(
        ['Tipo', 'Base de datos', 'Archivos', 'Tests'],
        [
            ['Unit tests', 'SQLite en memoria', 'test_audit_routes, test_edge_cases, test_security, test_statistics', '37'],
            ['Integration tests', 'PostgreSQL 16 real', 'test_integration_postgres', '12'],
            ['Total', '', '', '49'],
        ],
        col_widths=[1.3, 1.5, 2.7, 0.5]
    )
    
    add_section_title('8.2 Unit Tests (SQLite)', level=2)
    unit_points = [
        'Se ejecutan con DATABASE_URL=sqlite:///./test.db y APP_ENV=testing.',
        'Usan el tipo GUID que se adapta automáticamente a CHAR(36) para SQLite.',
        'No requieren PostgreSQL instalado.',
    ]
    for point in unit_points:
        p = doc.add_paragraph()
        p.add_run(f'• {point}')
    
    add_section_title('8.3 Integration Tests (PostgreSQL)', level=2)
    int_points = [
        'Requieren un servidor PostgreSQL 16 real corriendo.',
        'Se configuran con TEST_POSTGRES_URL=postgresql+asyncpg://...',
        'Ejecutan las migraciones Alembic antes de los tests.',
        'Validan funcionalidad real incluyendo el índice GIN full-text.',
    ]
    for point in int_points:
        p = doc.add_paragraph()
        p.add_run(f'• {point}')
    
    add_section_title('8.4 Pipeline CI/CD — GitHub Actions', level=2)
    add_body('**Archivo:** `.github/workflows/ci.yml` — 4 jobs configurados:')
    
    add_code_block(
        '┌─────────┐     ┌──────────────┐     ┌───────────────────┐     ┌──────────────┐\n'
        '│  lint   │────►│  test-unit   │────►│ test-integration  │────►│   docker     │\n'
        '│         │     │  (SQLite)    │     │  (PostgreSQL 16)  │     │  (build)     │\n'
        '└─────────┘     └──────────────┘     └───────────────────┘     └──────────────┘'
    )
    
    add_table(
        ['Job', 'Descripción', 'Depende de'],
        [
            ['lint', 'Verifica sintaxis Python y que los imports funcionen', '—'],
            ['test-unit', 'Ejecuta unit tests con SQLite', 'lint'],
            ['test-integration', 'Ejecuta integration tests con PostgreSQL 16 (service container)', 'lint'],
            ['docker', 'Construye imagen Docker multi-stage', 'test-unit + test-integration'],
        ],
        col_widths=[1.3, 3.2, 2.0]
    )
    
    add_note('El job docker solo se ejecuta en push a main (no en PRs ni en develop).')
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════════════
    # SECCIÓN 9: DEVOPS Y DESPLIEGUE
    # ══════════════════════════════════════════════════════════════════════
    
    add_section_title('9. DevOps y Despliegue')
    
    add_section_title('9.1 Dockerfile — Multi-Stage Build', level=2)
    
    add_table(
        ['Stage', 'Base', 'Propósito'],
        [
            ['builder', 'python:3.10-slim', 'Instala gcc, libpq-dev, compila dependencias'],
            ['runtime', 'python:3.10-slim', 'Solo libpq5 + curl + código + deps precompiladas'],
        ],
        col_widths=[1.0, 1.8, 3.7]
    )
    
    add_body('**Beneficios:** Imagen final más pequeña (sin compiladores), layer de dependencias cacheado, usuario no-root por seguridad.')
    
    add_body('**Healthcheck integrado:**')
    add_code_block(
        'HEALTHCHECK --interval=30s --timeout=5s --retries=3 \\\n'
        '    CMD curl -f http://localhost:8019/api/v1/audit/health || exit 1'
    )
    
    add_body('**CMD de ejecución** (ejecuta migraciones antes de iniciar):')
    add_code_block(
        'CMD ["sh", "-c", "python -m alembic upgrade head && \\\n'
        '     uvicorn app.main:app --host 0.0.0.0 --port 8019 \\\n'
        '     --workers 4 --loop uvloop --http httptools"]'
    )
    
    add_section_title('9.2 Docker Compose', level=2)
    
    add_table(
        ['Servicio', 'Imagen', 'Puerto', 'Descripción'],
        [
            ['db', 'postgres:16-alpine', '5432', 'PostgreSQL con healthcheck'],
            ['app', 'Build local', '8019', 'Microservicio ms-auditoria'],
        ],
        col_widths=[1.0, 1.5, 0.7, 3.3]
    )
    
    add_body('**Características:**')
    compose_points = [
        'depends_on con condition: service_healthy — la app espera a que PostgreSQL esté listo.',
        'restart: unless-stopped — reinicio automático del contenedor.',
        'Red erp-net (bridge) para comunicación con otros microservicios del ERP.',
        'Volumen pgdata para persistencia de datos entre reinicios.',
    ]
    for point in compose_points:
        p = doc.add_paragraph()
        p.add_run(f'• {point}')
    
    add_section_title('9.3 Variables de Entorno (Producción)', level=2)
    
    add_table(
        ['Variable', 'Valor', 'Descripción'],
        [
            ['DATABASE_URL', 'postgresql+psycopg2://postgres:***@db:5432/ms_auditoria', 'Conexión a PostgreSQL'],
            ['AES_SECRET_KEY', '(64 hex chars)', 'Clave AES-256'],
            ['APP_ENV', 'production', 'Entorno de ejecución'],
            ['APP_DEBUG', 'false', 'Sin modo debug'],
            ['LOG_LEVEL', 'INFO', 'Nivel de logging'],
            ['DB_POOL_SIZE', '20', 'Pool de conexiones'],
            ['DB_MAX_OVERFLOW', '40', 'Overflow del pool'],
            ['MS_AUTENTICACION_URL', 'http://ms-autenticacion:8001/...', 'URL ms-autenticación (red interna)'],
            ['MS_ROLES_URL', 'http://ms-roles-permisos:8002/...', 'URL ms-roles (red interna)'],
        ],
        col_widths=[1.8, 2.5, 2.2]
    )
    
    add_section_title('9.4 Lifecycle de la Aplicación', level=2)
    add_body('Gestionado con `asynccontextmanager` en `main.py`:')
    
    add_body('**Startup:**')
    startup_points = [
        'Log de inicio con entorno actual.',
        'Si APP_ENV == "development": crear tablas con Base.metadata.create_all() (sync engine).',
        'Iniciar scheduler de retención automática (retention_service.start()).',
    ]
    for i, point in enumerate(startup_points, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {point}')
    
    add_body('**Shutdown:**')
    shutdown_points = [
        'Detener scheduler de retención (retention_service.stop()).',
        'Cerrar pool async (async_engine.dispose()).',
        'Log de cierre.',
    ]
    for i, point in enumerate(shutdown_points, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {point}')
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════════════
    # SECCIÓN 10: JUSTIFICACIONES TÉCNICAS
    # ══════════════════════════════════════════════════════════════════════
    
    add_section_title('10. Justificaciones Técnicas')
    
    justifications = [
        ('10.1 ¿Por qué FastAPI?', [
            ('Async nativo', 'Soporte completo de async/await sin workarounds'),
            ('Rendimiento', 'Uno de los frameworks Python más rápidos (Starlette + Uvicorn)'),
            ('Documentación automática', 'Swagger UI (/docs) y ReDoc (/redoc) automáticos'),
            ('Validación integrada', 'Pydantic v2 para validación con tipos Python'),
            ('Dependency Injection', 'Sistema nativo de DI con Depends()'),
            ('Estándar OpenAPI', 'Compatible con generación de clientes'),
        ]),
        ('10.2 ¿Por qué SQLAlchemy 2.0 Async?', [
            ('Non-blocking I/O', 'Las consultas no bloquean el event loop'),
            ('Pool de conexiones', 'Gestión automática con pool_pre_ping, pool_recycle'),
            ('ORM maduro', 'Modelo de datos expresivo con tipos custom'),
            ('Compatibilidad', 'Funciona con PostgreSQL (asyncpg) y SQLite (aiosqlite)'),
            ('Select 2.0', 'Sintaxis select(Model).where(...) más explícita'),
        ]),
        ('10.3 ¿Por qué PostgreSQL 16?', [
            ('GIN index', 'Índice invertido para búsqueda full-text en español'),
            ('UUID nativo', 'Tipo UUID sin overhead de conversión'),
            ('JSONB', 'Soporte nativo si se necesita en el futuro'),
            ('Rendimiento', 'Mejoras en query planner y vacuuming'),
            ('Ecosystem', 'Driver async asyncpg con rendimiento superior'),
        ]),
        ('10.4 ¿Por qué Pydantic v2?', [
            ('Rendimiento', 'Core en Rust — hasta 50x más rápido que v1'),
            ('model_validate', 'Mapeo directo desde ORM con from_attributes=True'),
            ('computed_field', 'Campos calculados (ej: ASYNC_DATABASE_URL)'),
            ('ConfigDict', 'Configuración más limpia que class Config'),
            ('Integración FastAPI', 'Validación automática de request/response'),
        ]),
        ('10.5 ¿Por qué Repository Pattern?', [
            ('Testabilidad', 'Se puede mockear el repositorio en tests'),
            ('Separación', 'Consultas SQL aisladas de la lógica de negocio'),
            ('Mantenibilidad', 'Un solo lugar para modificar consultas'),
            ('Extensibilidad', 'Agregar filtros sin tocar el servicio'),
        ]),
        ('10.6 ¿Por qué Unit of Work como infraestructura disponible?', [
            ('Disponibilidad', 'Implementado en database/unit_of_work.py y disponible vía get_uow()'),
            ('Uso actual', 'Endpoints usan AsyncSession directa porque cada operación es una sola transacción simple'),
            ('Futuro', 'Queda como infraestructura para casos que requieran múltiples operaciones atómicas'),
        ]),
        ('10.7 ¿Por qué asyncio nativo para retención?', [
            ('Cero dependencias', 'No agrega librerías externas al proyecto'),
            ('Simplicidad', 'asyncio.create_task() + asyncio.sleep() es suficiente'),
            ('Integración', 'Se gestiona con el lifespan de FastAPI'),
            ('Caso simple', 'Solo una tarea diaria — no necesita scheduler completo'),
        ]),
        ('10.8 ¿Por qué JSON Structured Logging?', [
            ('Machine-readable', 'Parseable por ELK, Grafana, Loki'),
            ('Campos estándar', 'timestamp, level, service, message, module, function, line'),
            ('Extensible', 'Campos extra opcionales bajo key "extra"'),
            ('UTC', 'Timestamps en UTC ISO 8601 para consistencia'),
        ]),
        ('10.9 ¿Por qué el campo detalle NO se cifra automáticamente?', [
            ('Full-text search', 'El índice GIN requiere texto plano para funcionar'),
            ('Consultas', 'GET /logs?search_text=... necesita buscar dentro del detalle'),
            ('Naturaleza de datos', 'Los logs contienen metadata operativa, no datos sensibles'),
            ('Disponibilidad', 'Si se requiere, el microservicio emisor puede cifrar antes de enviar'),
        ]),
    ]
    
    for section_title, items in justifications:
        add_section_title(section_title, level=2)
        add_table(
            ['Razón', 'Detalle'],
            [[reason, detail] for reason, detail in items],
            col_widths=[1.5, 5.0]
        )
    
    add_body('**Ejemplo de log estructurado:**')
    add_code_block(
        '{\n'
        '  "timestamp": "2026-02-27T10:30:00.000000+00:00",\n'
        '  "level": "INFO",\n'
        '  "service": "ms-auditoria",\n'
        '  "message": "audit_log_created",\n'
        '  "module": "audit_service",\n'
        '  "function": "create_log",\n'
        '  "line": 45,\n'
        '  "extra": {\n'
        '    "audit_id": "a1b2c3d4-...",\n'
        '    "servicio": "ms-matriculas",\n'
        '    "endpoint": "/api/v1/matricula/inscribir"\n'
        '  }\n'
        '}'
    )
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════════════
    # APÉNDICE
    # ══════════════════════════════════════════════════════════════════════
    
    add_section_title('Apéndice: Configuración Completa')
    
    add_section_title('Variables de Entorno', level=2)
    
    add_table(
        ['Variable', 'Tipo', 'Default', 'Descripción'],
        [
            ['DATABASE_URL', 'str', 'postgresql+psycopg2://...', 'URL de conexión sync'],
            ['AES_SECRET_KEY', 'str', '(requerida)', 'Clave AES-256 (64 hex chars)'],
            ['API_KEY_HEADER', 'str', 'X-API-Key', 'Header de autenticación'],
            ['CORS_ORIGINS', 'str', 'localhost:3000,...', 'Orígenes CORS'],
            ['RATE_LIMIT_REQUESTS', 'int', '100', 'Máx requests por ventana'],
            ['RATE_LIMIT_WINDOW_SECONDS', 'int', '60', 'Ventana en segundos'],
            ['RETENTION_DAYS', 'int', '90', 'Días de retención'],
            ['RETENTION_CRON_HOUR', 'int', '3', 'Hora UTC de purga'],
            ['MS_AUTENTICACION_URL', 'str', 'http://localhost:8001/...', 'URL ms-autenticación'],
            ['MS_ROLES_URL', 'str', 'http://localhost:8002/...', 'URL ms-roles'],
            ['DB_POOL_SIZE', 'int', '10', 'Conexiones en pool'],
            ['DB_MAX_OVERFLOW', 'int', '20', 'Overflow del pool'],
            ['DB_POOL_RECYCLE', 'int', '3600', 'Reciclaje conexiones (seg)'],
            ['APP_HOST', 'str', '0.0.0.0', 'Host del servidor'],
            ['APP_PORT', 'int', '8019', 'Puerto del servidor'],
            ['APP_ENV', 'str', 'development', 'Entorno de ejecución'],
            ['APP_DEBUG', 'bool', 'False', 'Modo debug'],
            ['LOG_LEVEL', 'str', 'INFO', 'Nivel de logging'],
            ['DEFAULT_PAGE_SIZE', 'int', '20', 'Página default'],
            ['MAX_PAGE_SIZE', 'int', '100', 'Página máxima'],
        ],
        col_widths=[1.8, 0.5, 1.5, 2.5]
    )
    
    # ── Guardar ────────────────────────────────────────────────────────────
    output_path = os.path.join(os.path.dirname(__file__), 'Documento_Tecnico_Arquitectura.docx')
    doc.save(output_path)
    print(f'✅ Documento generado exitosamente: {output_path}')
    return output_path


if __name__ == '__main__':
    create_document()
